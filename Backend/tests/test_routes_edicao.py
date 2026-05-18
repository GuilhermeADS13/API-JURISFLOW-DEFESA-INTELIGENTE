"""Testes da rota POST /api/editar-contestacao (Fase 1 PR 2)."""

from __future__ import annotations

import asyncio
import base64
from io import BytesIO

import pytest
from docx import Document
from fastapi import HTTPException, Request

from App.routes import edicao as edicao_route
from App.services.n8n_service import N8NServiceError


def _docx_base_bytes() -> bytes:
    """Constroi um .docx base parecido com o caso de uso real."""
    doc = Document()
    doc.add_paragraph("CONTESTACAO")
    doc.add_paragraph("Processo numero 0000091-39.2026.5.06.0341")
    doc.add_paragraph("Reu: Janaina Pereira da Silva Matos, residente em Recife/PE.")
    doc.add_paragraph(
        "Por meio desta peca, a parte re Janaina Pereira da Silva Matos contesta os pedidos."
    )
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _payload_valido(**overrides) -> "edicao_route.EdicaoContestacao":
    docx_b64 = base64.b64encode(_docx_base_bytes()).decode("ascii")
    base = {
        "arquivo_base_conteudo_base64": docx_b64,
        "arquivo_base_nome": "Contestacao_VR.docx",
        "arquivo_base_mime_type": (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        "arquivo_base_tamanho_bytes": len(base64.b64decode(docx_b64)),
        "nome_novo": "Erica Cavalcante de Oliveira",
        "numero_processo_novo": "0000057-64.2026.5.06.0341",
        "valor_causa_novo": "27.598,41",
    }
    base.update(overrides)
    return edicao_route.EdicaoContestacao.model_validate(base)


def _fake_request(path: str = "/api/editar-contestacao") -> Request:
    """Request minimo para satisfazer o decorator @limiter.limit do slowapi."""
    return Request(
        scope={
            "type": "http",
            "method": "POST",
            "path": path,
            "headers": [],
            "query_string": b"",
            "client": ("127.0.0.1", 0),
        }
    )


# ── Validacao do schema (Pydantic) ──────────────────────────────────────────


def test_schema_recusa_payload_sem_nenhum_campo_novo():
    docx_b64 = base64.b64encode(_docx_base_bytes()).decode("ascii")
    with pytest.raises(Exception):  # ValidationError
        edicao_route.EdicaoContestacao.model_validate(
            {
                "arquivo_base_conteudo_base64": docx_b64,
                "arquivo_base_nome": "Contestacao.docx",
            }
        )


def test_schema_recusa_extensao_diferente_de_docx():
    docx_b64 = base64.b64encode(_docx_base_bytes()).decode("ascii")
    with pytest.raises(Exception):
        edicao_route.EdicaoContestacao.model_validate(
            {
                "arquivo_base_conteudo_base64": docx_b64,
                "arquivo_base_nome": "Contestacao.pdf",
                "nome_novo": "Erica Cavalcante",
            }
        )


def test_schema_recusa_base64_que_nao_eh_docx():
    """Base64 valido, mas conteudo nao tem magic bytes ZIP."""
    fake = base64.b64encode(b"not a real docx").decode("ascii")
    with pytest.raises(Exception):
        edicao_route.EdicaoContestacao.model_validate(
            {
                "arquivo_base_conteudo_base64": fake,
                "arquivo_base_nome": "Contestacao.docx",
                "nome_novo": "Erica Cavalcante",
            }
        )


def test_schema_recusa_numero_processo_fora_do_padrao_cnj():
    docx_b64 = base64.b64encode(_docx_base_bytes()).decode("ascii")
    with pytest.raises(Exception):
        edicao_route.EdicaoContestacao.model_validate(
            {
                "arquivo_base_conteudo_base64": docx_b64,
                "arquivo_base_nome": "Contestacao.docx",
                "numero_processo_novo": "12345",
            }
        )


def test_schema_recusa_valor_causa_em_formato_invalido():
    docx_b64 = base64.b64encode(_docx_base_bytes()).decode("ascii")
    with pytest.raises(Exception):
        edicao_route.EdicaoContestacao.model_validate(
            {
                "arquivo_base_conteudo_base64": docx_b64,
                "arquivo_base_nome": "Contestacao.docx",
                "valor_causa_novo": "27598.41",  # ponto como decimal — invalido
            }
        )


def test_schema_aceita_valor_com_prefixo_real():
    docx_b64 = base64.b64encode(_docx_base_bytes()).decode("ascii")
    payload = edicao_route.EdicaoContestacao.model_validate(
        {
            "arquivo_base_conteudo_base64": docx_b64,
            "arquivo_base_nome": "Contestacao.docx",
            "valor_causa_novo": "R$ 27.598,41",
        }
    )
    assert payload.valor_causa_novo == "R$ 27.598,41"


def test_schema_sanitiza_path_traversal_no_nome_do_arquivo():
    docx_b64 = base64.b64encode(_docx_base_bytes()).decode("ascii")
    payload = edicao_route.EdicaoContestacao.model_validate(
        {
            "arquivo_base_conteudo_base64": docx_b64,
            "arquivo_base_nome": "../../etc/passwd/Contestacao.docx",
            "nome_novo": "Erica Cavalcante",
        }
    )
    assert "/" not in payload.arquivo_base_nome
    assert payload.arquivo_base_nome == "Contestacao.docx"


# ── Rota: happy path com mock do n8n ────────────────────────────────────────


def test_happy_path_aplica_substituicoes_e_retorna_relatorio(monkeypatch):
    payload = _payload_valido()

    async def fake_enviar_n8n(dados):
        # O texto extraido tem 2 ocorrencias do nome e 1 do numero;
        # nao tem valor da causa.
        return {
            "substituicoes": [
                {
                    "campo": "nome",
                    "antigo": "Janaina Pereira da Silva Matos",
                    "novo": "Erica Cavalcante de Oliveira",
                    "ocorrencias_esperadas": 2,
                },
                {
                    "campo": "numero_processo",
                    "antigo": "0000091-39.2026.5.06.0341",
                    "novo": "0000057-64.2026.5.06.0341",
                    "ocorrencias_esperadas": 1,
                },
            ],
            "campos_ausentes": ["valor_causa"],
        }

    monkeypatch.setattr(edicao_route, "enviar_para_n8n_edicao", fake_enviar_n8n)

    resposta = asyncio.run(
        edicao_route.editar_contestacao(
            request=_fake_request(),
            payload=payload,
            usuario={"id": "USR-TESTE", "nome": "Ana", "email": "a@a.com"},
        )
    )

    assert resposta["status"] == "ok"
    assert resposta["arquivo_editado_nome"].endswith(".docx")
    assert resposta["campos_ausentes"] == ["valor_causa"]
    assert resposta["ocorrencias_aplicadas"] == {
        "Janaina Pereira da Silva Matos": 2,
        "0000091-39.2026.5.06.0341": 1,
    }
    relatorio = resposta["relatorio"]
    assert any("Erica Cavalcante de Oliveira" in linha for linha in relatorio)
    assert any("0000057-64.2026.5.06.0341" in linha for linha in relatorio)
    assert any(
        "valor da causa" in linha.lower() and "27.598,41" in linha
        for linha in relatorio
    )

    # O .docx editado deve abrir e conter o texto novo, sem o antigo.
    docx_editado = base64.b64decode(resposta["arquivo_editado_base64"])
    doc = Document(BytesIO(docx_editado))
    texto_total = "\n".join(p.text for p in doc.paragraphs)
    assert "Erica Cavalcante de Oliveira" in texto_total
    assert "Janaina Pereira da Silva Matos" not in texto_total
    assert "0000057-64.2026.5.06.0341" in texto_total


def test_divergencia_de_ocorrencias_aborta_com_422(monkeypatch):
    """Agente retorna ocorrencias_esperadas que nao bate com o documento."""
    payload = _payload_valido()

    async def fake_enviar_n8n(dados):
        return {
            "substituicoes": [
                {
                    "campo": "nome",
                    "antigo": "Janaina Pereira da Silva Matos",
                    "novo": "Erica Cavalcante de Oliveira",
                    "ocorrencias_esperadas": 5,  # texto so tem 2!
                },
            ],
            "campos_ausentes": [],
        }

    monkeypatch.setattr(edicao_route, "enviar_para_n8n_edicao", fake_enviar_n8n)

    with pytest.raises(HTTPException) as exc:
        asyncio.run(
            edicao_route.editar_contestacao(
                request=_fake_request(),
                payload=payload,
                usuario={"id": "USR-TESTE", "nome": "Ana", "email": "a@a.com"},
            )
        )
    assert exc.value.status_code == 422
    assert "substitui" in str(exc.value.detail).lower()


def test_workflow_indisponivel_retorna_503(monkeypatch):
    payload = _payload_valido()

    async def fake_enviar_n8n(dados):
        raise N8NServiceError("connection refused")

    monkeypatch.setattr(edicao_route, "enviar_para_n8n_edicao", fake_enviar_n8n)

    with pytest.raises(HTTPException) as exc:
        asyncio.run(
            edicao_route.editar_contestacao(
                request=_fake_request(),
                payload=payload,
                usuario={"id": "USR-TESTE", "nome": "Ana", "email": "a@a.com"},
            )
        )
    assert exc.value.status_code == 503
    # Mensagem generica, nao vaza detalhe do connection refused.
    assert "connection" not in str(exc.value.detail).lower()


def test_resposta_n8n_fora_do_schema_retorna_502(monkeypatch):
    payload = _payload_valido()

    async def fake_enviar_n8n(dados):
        # Sem `substituicoes` nem `campos_ausentes` no formato esperado.
        return {"foo": "bar", "substituicoes": [{"sem_campos_esperados": True}]}

    monkeypatch.setattr(edicao_route, "enviar_para_n8n_edicao", fake_enviar_n8n)

    with pytest.raises(HTTPException) as exc:
        asyncio.run(
            edicao_route.editar_contestacao(
                request=_fake_request(),
                payload=payload,
                usuario={"id": "USR-TESTE", "nome": "Ana", "email": "a@a.com"},
            )
        )
    assert exc.value.status_code == 502


def test_resposta_n8n_nao_dict_retorna_502(monkeypatch):
    payload = _payload_valido()

    async def fake_enviar_n8n(dados):
        return ["lista", "em", "vez", "de", "dict"]

    monkeypatch.setattr(edicao_route, "enviar_para_n8n_edicao", fake_enviar_n8n)

    with pytest.raises(HTTPException) as exc:
        asyncio.run(
            edicao_route.editar_contestacao(
                request=_fake_request(),
                payload=payload,
                usuario={"id": "USR-TESTE", "nome": "Ana", "email": "a@a.com"},
            )
        )
    assert exc.value.status_code == 502


def test_todos_campos_pedidos_ausentes_devolve_arquivo_original(monkeypatch):
    payload = _payload_valido()

    async def fake_enviar_n8n(dados):
        return {
            "substituicoes": [],
            "campos_ausentes": ["nome", "numero_processo", "valor_causa"],
        }

    monkeypatch.setattr(edicao_route, "enviar_para_n8n_edicao", fake_enviar_n8n)

    resposta = asyncio.run(
        edicao_route.editar_contestacao(
            request=_fake_request(),
            payload=payload,
            usuario={"id": "USR-TESTE", "nome": "Ana", "email": "a@a.com"},
        )
    )

    assert resposta["status"] == "ok"
    assert resposta["ocorrencias_aplicadas"] == {}
    # 3 campos ausentes -> 3 bullets explicativos no relatorio.
    assert len(resposta["relatorio"]) == 3
    assert all("Nao havia" in linha for linha in resposta["relatorio"])


def test_nome_arquivo_saida_usa_numero_processo_novo():
    payload = _payload_valido()
    nome = edicao_route._montar_nome_saida(payload)
    assert nome == "contestacao_editada_0000057_64_2026_5_06_0341.docx"


def test_nome_arquivo_saida_padrao_quando_sem_numero():
    docx_b64 = base64.b64encode(_docx_base_bytes()).decode("ascii")
    payload = edicao_route.EdicaoContestacao.model_validate(
        {
            "arquivo_base_conteudo_base64": docx_b64,
            "arquivo_base_nome": "Contestacao.docx",
            "nome_novo": "Erica Cavalcante",
        }
    )
    nome = edicao_route._montar_nome_saida(payload)
    assert nome == "contestacao_editada.docx"
