"""Testes da rota POST /api/contestar-por-peticao (Guia Tecnico v2 - PR2)."""

from __future__ import annotations

import asyncio
import base64
from io import BytesIO

import pytest
from docx import Document
from fastapi import HTTPException, Request

from App.routes import contestacao_peticao as peticao_route
from App.services.n8n_service import N8NServiceError


def _docx_peticao_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("PETICAO INICIAL")
    doc.add_paragraph(
        "Reclamante: Joao da Silva, brasileiro, residente em Recife/PE."
    )
    doc.add_paragraph(
        "Reclamada: Empresa XYZ LTDA. Pleiteia horas extras nao pagas e adicional noturno."
    )
    doc.add_paragraph("Valor da causa: R$ 27.598,41.")
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _modelo_base_docx_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("Modelo do escritorio para {{ tipo_acao }}.")
    doc.add_paragraph("Autor: {{ autor }}. Reu: {{ reu }}.")
    doc.add_paragraph("Tese central: {{ tese_central }}")
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _payload_valido(**overrides) -> "peticao_route.ContestacaoPorPeticao":
    base = {
        "arquivo_peticao_base64": base64.b64encode(_docx_peticao_bytes()).decode("ascii"),
        "arquivo_peticao_nome": "peticao.docx",
        "arquivo_peticao_mime_type": (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        "tipo_acao_hint": "Trabalhista — Horas Extras",
        "pontos_contestante": "Atacar prescricao bienal e ausencia de prova das horas extras.",
    }
    base.update(overrides)
    return peticao_route.ContestacaoPorPeticao.model_validate(base)


def _fake_request() -> Request:
    return Request(
        scope={
            "type": "http",
            "method": "POST",
            "path": "/api/contestar-por-peticao",
            "headers": [],
            "query_string": b"",
            "client": ("127.0.0.1", 0),
        }
    )


def _resposta_n8n_ok() -> dict:
    return {
        "status": "ok",
        "dados_extraidos": {
            "numero_processo": "0001234-56.2026.5.06.0341",
            "autor": "Joao da Silva",
            "reu": "Empresa XYZ LTDA",
            "tipo_acao": "Trabalhista — Horas Extras",
            "vara": "1a Vara do Trabalho de Recife",
            "fatos_resumo": "Reclamante alega horas extras nao pagas.",
            "pedidos": ["Horas extras", "Adicional noturno"],
            "valores": {"total_estimado": "R$ 27.598,41"},
            "argumentos_autor": ["Jornada extensa", "Falta de cartoes ponto"],
            "pontos_vulneraveis": ["Ausencia de prova testemunhal"],
            "confianca": 0.92,
        },
        "minuta": {
            "tese_central": "Improcedencia total dos pedidos.",
            "merito": "Os pedidos do autor carecem de fundamento.",
            "fundamentos": "Art. 818 CLT, art. 373 CPC.",
            "pedidos": "Improcedencia integral.",
            "impugnacao_pedidos": {
                "Horas extras": "Nao houve labor extraordinario.",
                "Adicional noturno": "Nao se aplica ao caso.",
            },
            "preliminares": None,
            "observacoes": "Revisar antes de protocolar.",
        },
        "engine_ia": {"provider": "claude", "model": "claude-sonnet-4-6"},
    }


# ── Validacao do schema ─────────────────────────────────────────────────────


def test_schema_recusa_extensao_invalida():
    with pytest.raises(Exception):
        peticao_route.ContestacaoPorPeticao.model_validate({
            "arquivo_peticao_base64": base64.b64encode(_docx_peticao_bytes()).decode("ascii"),
            "arquivo_peticao_nome": "peticao.txt",
        })


def test_schema_recusa_arquivo_muito_grande():
    grande = b"%PDF-1.4\n" + b"a" * (21 * 1024 * 1024)
    with pytest.raises(Exception):
        peticao_route.ContestacaoPorPeticao.model_validate({
            "arquivo_peticao_base64": base64.b64encode(grande).decode("ascii"),
            "arquivo_peticao_nome": "peticao.pdf",
        })


# ── Happy path: rota gera DOCX a partir da resposta do n8n ──────────────────


def test_happy_path_sem_modelo_base(monkeypatch):
    payload = _payload_valido()

    async def fake_n8n(dados):
        # Confirma que o backend ja extraiu o texto antes de chamar o n8n.
        assert "texto_peticao" in dados
        assert "Reclamante" in dados["texto_peticao"]
        assert dados["tipo_acao_hint"] == "Trabalhista — Horas Extras"
        return _resposta_n8n_ok()

    monkeypatch.setattr(peticao_route, "enviar_para_n8n_peticao", fake_n8n)
    monkeypatch.setattr(peticao_route, "save_contestacao", lambda **kw: 42)

    resposta = asyncio.run(
        peticao_route.contestar_por_peticao(
            request=_fake_request(),
            payload=payload,
            usuario={"id": "USR-TESTE", "nome": "Ana", "email": "a@a.com"},
        )
    )

    assert resposta["status"] == "ok"
    assert resposta["contestacao_id"] == 42
    assert resposta["dados_extraidos"]["autor"] == "Joao da Silva"
    assert resposta["arquivo_editado_nome"].endswith(".docx")

    docx_bytes = base64.b64decode(resposta["arquivo_editado_base64"])
    assert docx_bytes.startswith(b"PK\x03\x04")  # ZIP/OpenXML
    doc = Document(BytesIO(docx_bytes))
    texto = "\n".join(p.text for p in doc.paragraphs)
    assert "CONTESTACAO" in texto
    assert "Improcedencia total" in texto


def test_happy_path_com_modelo_base_chama_docxtpl(monkeypatch):
    payload = _payload_valido(
        modelo_base_base64=base64.b64encode(_modelo_base_docx_bytes()).decode("ascii"),
        modelo_base_nome="modelo_escritorio.docx",
    )

    async def fake_n8n(dados):
        assert dados["modelo_base_texto"]  # backend extraiu o texto do modelo
        return _resposta_n8n_ok()

    monkeypatch.setattr(peticao_route, "enviar_para_n8n_peticao", fake_n8n)
    monkeypatch.setattr(peticao_route, "save_contestacao", lambda **kw: 99)

    resposta = asyncio.run(
        peticao_route.contestar_por_peticao(
            request=_fake_request(),
            payload=payload,
            usuario={"id": "USR-TESTE", "nome": "Ana", "email": "a@a.com"},
        )
    )

    assert resposta["status"] == "ok"
    docx_bytes = base64.b64decode(resposta["arquivo_editado_base64"])
    doc = Document(BytesIO(docx_bytes))
    texto = "\n".join(p.text for p in doc.paragraphs)
    # Placeholders devem ter sido substituidos pelo docxtpl.
    assert "Joao da Silva" in texto
    assert "Empresa XYZ LTDA" in texto
    assert "{{" not in texto


# ── Falhas ───────────────────────────────────────────────────────────────────


def test_n8n_indisponivel_retorna_502(monkeypatch):
    payload = _payload_valido()

    async def fake_n8n(dados):
        raise N8NServiceError("connection refused")

    monkeypatch.setattr(peticao_route, "enviar_para_n8n_peticao", fake_n8n)

    with pytest.raises(HTTPException) as exc:
        asyncio.run(
            peticao_route.contestar_por_peticao(
                request=_fake_request(),
                payload=payload,
                usuario={"id": "USR-TESTE", "nome": "Ana", "email": "a@a.com"},
            )
        )
    assert exc.value.status_code == 502
    assert "connection" not in str(exc.value.detail).lower()


def test_n8n_resposta_sem_minuta_retorna_502(monkeypatch):
    payload = _payload_valido()

    async def fake_n8n(dados):
        return {"status": "ok", "dados_extraidos": {}}  # sem minuta

    monkeypatch.setattr(peticao_route, "enviar_para_n8n_peticao", fake_n8n)

    with pytest.raises(HTTPException) as exc:
        asyncio.run(
            peticao_route.contestar_por_peticao(
                request=_fake_request(),
                payload=payload,
                usuario={"id": "USR-TESTE", "nome": "Ana", "email": "a@a.com"},
            )
        )
    assert exc.value.status_code == 502


def test_n8n_resposta_nao_dict_retorna_502(monkeypatch):
    payload = _payload_valido()

    async def fake_n8n(dados):
        return ["lista", "em", "vez", "de", "dict"]

    monkeypatch.setattr(peticao_route, "enviar_para_n8n_peticao", fake_n8n)

    with pytest.raises(HTTPException) as exc:
        asyncio.run(
            peticao_route.contestar_por_peticao(
                request=_fake_request(),
                payload=payload,
                usuario={"id": "USR-TESTE", "nome": "Ana", "email": "a@a.com"},
            )
        )
    assert exc.value.status_code == 502


def test_extracao_falha_retorna_422(monkeypatch):
    """Bytes que comecam com %PDF mas nao sao um PDF valido."""
    pdf_quebrado = b"%PDF-1.4\n" + b"corrompido " * 30
    payload = peticao_route.ContestacaoPorPeticao.model_validate({
        "arquivo_peticao_base64": base64.b64encode(pdf_quebrado).decode("ascii"),
        "arquivo_peticao_nome": "peticao.pdf",
    })

    # Forca o extractor a falhar.
    def fake_extract(peticao_bytes, peticao_nome, anexos=None):
        from App.services.peticao_extractor import ExtracaoError
        raise ExtracaoError("Nao foi possivel extrair texto legivel.")

    monkeypatch.setattr(peticao_route, "extrair_e_consolidar_textos", fake_extract)

    with pytest.raises(HTTPException) as exc:
        asyncio.run(
            peticao_route.contestar_por_peticao(
                request=_fake_request(),
                payload=payload,
                usuario={"id": "USR-TESTE", "nome": "Ana", "email": "a@a.com"},
            )
        )
    assert exc.value.status_code == 422


def test_origem_peticao_persistida(monkeypatch):
    payload = _payload_valido()
    save_calls = {}

    async def fake_n8n(dados):
        return _resposta_n8n_ok()

    def fake_save(**kw):
        save_calls.update(kw)
        return 7

    monkeypatch.setattr(peticao_route, "enviar_para_n8n_peticao", fake_n8n)
    monkeypatch.setattr(peticao_route, "save_contestacao", fake_save)

    asyncio.run(
        peticao_route.contestar_por_peticao(
            request=_fake_request(),
            payload=payload,
            usuario={"id": "USR-TESTE", "nome": "Ana", "email": "a@a.com"},
        )
    )

    assert save_calls["origem"] == "peticao"
    # Nao persistimos o conteudo da peticao inteira no banco.
    assert save_calls["payload"]["arquivo_base_conteudo_base64"] == ""


# ── PR5 HiL: Human-in-the-Loop ─────────────────────────────────────────────


def test_confianca_baixa_marca_revisao_humana(monkeypatch):
    """Quando confianca < 0.7 a rota retorna requer_revisao_humana sem DOCX."""
    payload = _payload_valido()
    save_calls = {}

    resposta_baixa_confianca = _resposta_n8n_ok()
    resposta_baixa_confianca["dados_extraidos"]["confianca"] = 0.42

    async def fake_n8n(dados):
        return resposta_baixa_confianca

    def fake_save(**kw):
        save_calls.update(kw)
        return 1234

    monkeypatch.setattr(peticao_route, "enviar_para_n8n_peticao", fake_n8n)
    monkeypatch.setattr(peticao_route, "save_contestacao", fake_save)

    resposta = asyncio.run(
        peticao_route.contestar_por_peticao(
            request=_fake_request(),
            payload=payload,
            usuario={"id": "USR-TESTE", "nome": "Ana", "email": "a@a.com"},
        )
    )

    assert resposta["status"] == "requer_revisao_humana"
    assert resposta["requer_revisao_humana"] is True
    assert resposta["dados_confianca"] == 0.42
    assert "arquivo_editado_base64" not in resposta  # nao gera DOCX

    # Persistido com flag e confianca
    assert save_calls["requer_revisao_humana"] is True
    assert save_calls["dados_confianca"] == 0.42
    assert save_calls["status"] == "requer_revisao_humana"
    # Minuta original preservada para golden dataset (PR5 Observabilidade)
    assert save_calls["minuta_json_original"] is not None


def test_confianca_alta_gera_docx_e_persiste_minuta_original(monkeypatch):
    """Confianca >= 0.7 → fluxo normal com DOCX e minuta_json_original gravada."""
    payload = _payload_valido()
    save_calls = {}

    async def fake_n8n(dados):
        return _resposta_n8n_ok()  # confianca padrao 0.92

    def fake_save(**kw):
        save_calls.update(kw)
        return 99

    monkeypatch.setattr(peticao_route, "enviar_para_n8n_peticao", fake_n8n)
    monkeypatch.setattr(peticao_route, "save_contestacao", fake_save)

    resposta = asyncio.run(
        peticao_route.contestar_por_peticao(
            request=_fake_request(),
            payload=payload,
            usuario={"id": "USR-TESTE", "nome": "Ana", "email": "a@a.com"},
        )
    )

    assert resposta["status"] == "ok"
    assert resposta["requer_revisao_humana"] is False
    assert resposta["dados_confianca"] == 0.92
    assert resposta["arquivo_editado_base64"]

    assert save_calls["requer_revisao_humana"] is False
    assert save_calls["minuta_json_original"]["tese_central"] == "Improcedencia total dos pedidos."


def test_confirmar_extracao_404_quando_nao_pertence(monkeypatch):
    monkeypatch.setattr(peticao_route, "get_contestacao", lambda cid, uid: None)

    payload = peticao_route.ConfirmacaoExtracao.model_validate({
        "dados_extraidos": {"autor": "X", "tipo_acao": "Y", "pedidos": []},
    })

    with pytest.raises(HTTPException) as exc:
        asyncio.run(
            peticao_route.confirmar_extracao(
                request=_fake_request(),
                contestacao_id=999,
                payload=payload,
                usuario={"id": "U", "nome": "n", "email": "e@e"},
            )
        )
    assert exc.value.status_code == 404


def test_confirmar_extracao_409_quando_nao_em_revisao(monkeypatch):
    monkeypatch.setattr(
        peticao_route,
        "get_contestacao",
        lambda cid, uid: {"id": cid, "requer_revisao_humana": False, "status": "ok"},
    )

    payload = peticao_route.ConfirmacaoExtracao.model_validate({
        "dados_extraidos": {"autor": "X", "tipo_acao": "Y", "pedidos": []},
    })

    with pytest.raises(HTTPException) as exc:
        asyncio.run(
            peticao_route.confirmar_extracao(
                request=_fake_request(),
                contestacao_id=10,
                payload=payload,
                usuario={"id": "U", "nome": "n", "email": "e@e"},
            )
        )
    assert exc.value.status_code == 409


def test_confirmar_extracao_envia_flag_pre_validados(monkeypatch):
    """Workflow n8n recebe `dados_extraidos_pre_validados` no payload."""
    payload_n8n_capturado = {}

    async def fake_n8n(dados):
        payload_n8n_capturado.update(dados)
        return _resposta_n8n_ok()

    monkeypatch.setattr(
        peticao_route,
        "get_contestacao",
        lambda cid, uid: {
            "id": cid, "requer_revisao_humana": True, "status": "requer_revisao_humana",
        },
    )
    monkeypatch.setattr(peticao_route, "enviar_para_n8n_peticao", fake_n8n)
    monkeypatch.setattr(
        peticao_route,
        "atualizar_contestacao_pos_revisao",
        lambda **kw: True,
    )

    payload = peticao_route.ConfirmacaoExtracao.model_validate({
        "dados_extraidos": {
            "autor": "Joao Corrigido",
            "tipo_acao": "Trabalhista",
            "reu": "Empresa X",
            "fatos_resumo": "Fatos corrigidos pelo humano.",
            "pedidos": ["pedido 1", "pedido 2"],
        },
    })

    resposta = asyncio.run(
        peticao_route.confirmar_extracao(
            request=_fake_request(),
            contestacao_id=42,
            payload=payload,
            usuario={"id": "U", "nome": "n", "email": "e@e"},
        )
    )

    assert resposta["status"] == "ok"
    assert resposta["requer_revisao_humana"] is False
    assert resposta["arquivo_editado_base64"]
    # Flag de bypass do extrator chega ao workflow
    assert "dados_extraidos_pre_validados" in payload_n8n_capturado
    assert payload_n8n_capturado["dados_extraidos_pre_validados"]["autor"] == "Joao Corrigido"


# ── PR5 Multi-docs: anexos ─────────────────────────────────────────────────


def test_payload_aceita_lista_anexos():
    payload = peticao_route.ContestacaoPorPeticao.model_validate({
        "arquivo_peticao_base64": base64.b64encode(_docx_peticao_bytes()).decode("ascii"),
        "arquivo_peticao_nome": "peticao.docx",
        "arquivos_anexos": [
            {
                "base64": base64.b64encode(_docx_peticao_bytes()).decode("ascii"),
                "nome": "anexo1.docx",
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            },
        ],
    })
    assert len(payload.arquivos_anexos) == 1
    assert payload.arquivos_anexos[0].nome == "anexo1.docx"


def test_payload_recusa_mais_de_5_anexos():
    docx_b64 = base64.b64encode(_docx_peticao_bytes()).decode("ascii")
    with pytest.raises(Exception):  # ValidationError
        peticao_route.ContestacaoPorPeticao.model_validate({
            "arquivo_peticao_base64": docx_b64,
            "arquivo_peticao_nome": "peticao.docx",
            "arquivos_anexos": [
                {"base64": docx_b64, "nome": f"a{i}.docx"} for i in range(6)
            ],
        })


def test_consolidacao_inclui_separadores_de_anexo(monkeypatch):
    """O texto enviado ao n8n contem cabecalhos === ANEXO N === por anexo."""
    docx_b64 = base64.b64encode(_docx_peticao_bytes()).decode("ascii")
    payload = peticao_route.ContestacaoPorPeticao.model_validate({
        "arquivo_peticao_base64": docx_b64,
        "arquivo_peticao_nome": "peticao.docx",
        "arquivos_anexos": [
            {"base64": docx_b64, "nome": "contrato.docx"},
        ],
    })

    payload_n8n_capturado = {}

    async def fake_n8n(dados):
        payload_n8n_capturado.update(dados)
        return _resposta_n8n_ok()

    monkeypatch.setattr(peticao_route, "enviar_para_n8n_peticao", fake_n8n)
    monkeypatch.setattr(peticao_route, "save_contestacao", lambda **kw: 1)

    asyncio.run(
        peticao_route.contestar_por_peticao(
            request=_fake_request(),
            payload=payload,
            usuario={"id": "U", "nome": "n", "email": "e@e"},
        )
    )

    texto = payload_n8n_capturado["texto_peticao"]
    assert "=== ANEXO 1 (contrato.docx) ===" in texto


# ── PR5 Observabilidade: PATCH minuta editada ─────────────────────────────


def test_patch_minuta_atualiza_editada(monkeypatch):
    """PATCH /contestacoes/{id}/minuta grava em minuta_json_editada."""
    save_calls = {}

    monkeypatch.setattr(
        peticao_route,
        "get_contestacao",
        lambda cid, uid: {
            "id": cid,
            "requer_revisao_humana": False,
            "minuta_json_original": {"tese_central": "Tese original IA", "merito": "Merito IA"},
            "minuta_json_editada": None,
        },
    )

    def fake_salvar(**kw):
        save_calls.update(kw)
        return True

    monkeypatch.setattr(peticao_route, "salvar_minuta_editada", fake_salvar)

    payload = peticao_route.MinutaEditada.model_validate({
        "tese_central": "Tese reescrita pelo advogado",
        "merito": "Merito ampliado pelo advogado",
    })

    resposta = asyncio.run(
        peticao_route.atualizar_minuta_editada(
            request=_fake_request(),
            contestacao_id=42,
            payload=payload,
            usuario={"id": "U", "nome": "n", "email": "e@e"},
        )
    )

    assert resposta["status"] == "ok"
    assert resposta["contestacao_id"] == 42
    assert "diff_resumo" in resposta
    assert resposta["diff_resumo"]["total_secoes_alteradas"] >= 1

    # Persistiu so as secoes editadas, original intacta.
    assert save_calls["minuta_editada"]["tese_central"] == "Tese reescrita pelo advogado"
    assert save_calls["minuta_editada"]["merito"] == "Merito ampliado pelo advogado"


def test_patch_minuta_404_quando_nao_pertence(monkeypatch):
    monkeypatch.setattr(peticao_route, "get_contestacao", lambda cid, uid: None)

    payload = peticao_route.MinutaEditada.model_validate({"tese_central": "x"})
    with pytest.raises(HTTPException) as exc:
        asyncio.run(
            peticao_route.atualizar_minuta_editada(
                request=_fake_request(),
                contestacao_id=999,
                payload=payload,
                usuario={"id": "U", "nome": "n", "email": "e@e"},
            )
        )
    assert exc.value.status_code == 404


def test_patch_minuta_recusa_payload_vazio():
    """Pelo menos 1 campo obrigatorio (validador do MinutaEditada)."""
    with pytest.raises(Exception):  # ValidationError
        peticao_route.MinutaEditada.model_validate({})
