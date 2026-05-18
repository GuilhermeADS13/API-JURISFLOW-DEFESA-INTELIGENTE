"""Testes do editor ciruurgico de .docx.

As fixtures sao construidas em memoria via python-docx para nao depender de
arquivos binarios commitados no repo.
"""

from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document

from App.services.docx_editor import (
    SubstituicaoError,
    aplicar_substituicoes,
    extrair_texto,
)


def _docx_simples(texto: str) -> bytes:
    """Cria um .docx com um paragrafo unico contendo `texto`."""
    doc = Document()
    doc.add_paragraph(texto)
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _docx_com_runs_fragmentados(*trechos: str) -> bytes:
    """Cria um paragrafo com varios runs, um por trecho passado.

    Util para simular um .docx onde a string que queremos substituir cruza
    fronteiras de <w:r> (cenario tipico de docs editados manualmente).
    """
    doc = Document()
    paragraph = doc.add_paragraph()
    for trecho in trechos:
        paragraph.add_run(trecho)
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _docx_com_tabela(linhas: list[list[str]]) -> bytes:
    doc = Document()
    table = doc.add_table(rows=len(linhas), cols=len(linhas[0]))
    for i, linha in enumerate(linhas):
        for j, valor in enumerate(linha):
            table.rows[i].cells[j].text = valor
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _ler_paragrafos(docx_bytes: bytes) -> list[str]:
    doc = Document(BytesIO(docx_bytes))
    return [p.text for p in doc.paragraphs]


def _ler_tabela(docx_bytes: bytes) -> list[list[str]]:
    doc = Document(BytesIO(docx_bytes))
    return [[cell.text for cell in row.cells] for row in doc.tables[0].rows]


def test_substituicao_simples_em_run_unico():
    docx = _docx_simples("Reu: Janaina Pereira da Silva Matos.")
    pares = [
        {
            "antigo": "Janaina Pereira da Silva Matos",
            "novo": "Erica Cavalcante de Oliveira",
        }
    ]

    novo_bytes, ocorrencias = aplicar_substituicoes(docx, pares)

    assert ocorrencias == {"Janaina Pereira da Silva Matos": 1}
    assert _ler_paragrafos(novo_bytes) == ["Reu: Erica Cavalcante de Oliveira."]


def test_multiplas_ocorrencias_no_mesmo_paragrafo():
    docx = _docx_simples(
        "Processo 0000091-39.2026.5.06.0341 vinculado ao caso 0000091-39.2026.5.06.0341."
    )
    pares = [
        {"antigo": "0000091-39.2026.5.06.0341", "novo": "0000057-64.2026.5.06.0341"}
    ]

    novo_bytes, ocorrencias = aplicar_substituicoes(docx, pares)

    assert ocorrencias == {"0000091-39.2026.5.06.0341": 2}
    assert _ler_paragrafos(novo_bytes) == [
        "Processo 0000057-64.2026.5.06.0341 vinculado ao caso 0000057-64.2026.5.06.0341."
    ]


def test_substituicao_em_runs_fragmentados():
    """Texto que cruza <w:r> (ex: parte do nome em negrito) deve ser substituido."""
    docx = _docx_com_runs_fragmentados(
        "Reu: Janaina ", "Pereira ", "da Silva ", "Matos."
    )
    pares = [
        {
            "antigo": "Janaina Pereira da Silva Matos",
            "novo": "Erica Cavalcante de Oliveira",
        }
    ]

    novo_bytes, ocorrencias = aplicar_substituicoes(docx, pares)

    assert ocorrencias == {"Janaina Pereira da Silva Matos": 1}
    assert _ler_paragrafos(novo_bytes) == ["Reu: Erica Cavalcante de Oliveira."]


def test_ocorrencia_inexistente_retorna_zero():
    docx = _docx_simples("Nada para substituir aqui.")
    pares = [{"antigo": "Joao", "novo": "Maria"}]

    novo_bytes, ocorrencias = aplicar_substituicoes(docx, pares)

    assert ocorrencias == {"Joao": 0}
    assert _ler_paragrafos(novo_bytes) == ["Nada para substituir aqui."]


def test_substituicao_dentro_de_tabela():
    docx = _docx_com_tabela(
        [
            ["Campo", "Valor"],
            ["Reclamado", "Janaina Pereira"],
            ["Valor da causa", "R$ 10.000,00"],
        ]
    )
    pares = [
        {"antigo": "Janaina Pereira", "novo": "Erica Cavalcante"},
        {"antigo": "R$ 10.000,00", "novo": "R$ 27.598,41"},
    ]

    novo_bytes, ocorrencias = aplicar_substituicoes(docx, pares)

    assert ocorrencias == {"Janaina Pereira": 1, "R$ 10.000,00": 1}
    tabela = _ler_tabela(novo_bytes)
    assert tabela[1] == ["Reclamado", "Erica Cavalcante"]
    assert tabela[2] == ["Valor da causa", "R$ 27.598,41"]


def test_pares_multiplos_em_uma_passada():
    docx = _docx_simples(
        "Contestacao do processo 0000091-39.2026.5.06.0341, "
        "ajuizado em face de Janaina Pereira, valor da causa R$ 10.000,00."
    )
    pares = [
        {"antigo": "Janaina Pereira", "novo": "Erica Cavalcante"},
        {"antigo": "0000091-39.2026.5.06.0341", "novo": "0000057-64.2026.5.06.0341"},
        {"antigo": "R$ 10.000,00", "novo": "R$ 27.598,41"},
    ]

    novo_bytes, ocorrencias = aplicar_substituicoes(docx, pares)

    assert ocorrencias == {
        "Janaina Pereira": 1,
        "0000091-39.2026.5.06.0341": 1,
        "R$ 10.000,00": 1,
    }
    assert _ler_paragrafos(novo_bytes) == [
        "Contestacao do processo 0000057-64.2026.5.06.0341, "
        "ajuizado em face de Erica Cavalcante, valor da causa R$ 27.598,41."
    ]


def test_par_com_antigo_vazio_eh_ignorado():
    docx = _docx_simples("Texto qualquer.")
    pares = [
        {"antigo": "", "novo": "qualquer coisa"},
        {"antigo": "Texto", "novo": "Frase"},
    ]

    novo_bytes, ocorrencias = aplicar_substituicoes(docx, pares)

    # Antigo vazio nao entra no resultado de ocorrencias.
    assert ocorrencias == {"Texto": 1}
    assert _ler_paragrafos(novo_bytes) == ["Frase qualquer."]


def test_preserva_estilo_quando_substituicao_cabe_no_run():
    """Quando o `antigo` esta inteiro num run com bold, o run editado deve manter bold."""
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("Reu: ")
    bold_run = paragraph.add_run("Janaina Pereira")
    bold_run.bold = True
    paragraph.add_run(" residente em SP.")
    out = BytesIO()
    doc.save(out)

    pares = [{"antigo": "Janaina Pereira", "novo": "Erica Cavalcante"}]
    novo_bytes, ocorrencias = aplicar_substituicoes(out.getvalue(), pares)

    assert ocorrencias == {"Janaina Pereira": 1}
    novo_doc = Document(BytesIO(novo_bytes))
    runs = novo_doc.paragraphs[0].runs
    assert runs[0].text == "Reu: "
    assert runs[1].text == "Erica Cavalcante"
    assert runs[1].bold is True
    assert runs[2].text == " residente em SP."


def test_bytes_vazios_levanta_erro():
    with pytest.raises(SubstituicaoError):
        aplicar_substituicoes(b"", [{"antigo": "x", "novo": "y"}])


def test_bytes_invalidos_levantam_erro():
    with pytest.raises(SubstituicaoError):
        aplicar_substituicoes(b"nao eh um docx valido", [{"antigo": "x", "novo": "y"}])


def test_resultado_eh_docx_valido_e_relegivel():
    """Smoke test: o output bate como .docx legivel pelo python-docx."""
    docx = _docx_simples("Original: Joao da Silva.")
    pares = [{"antigo": "Joao da Silva", "novo": "Maria Souza"}]

    novo_bytes, _ = aplicar_substituicoes(docx, pares)

    # Deve abrir sem erro e ter o texto novo.
    doc = Document(BytesIO(novo_bytes))
    assert doc.paragraphs[0].text == "Original: Maria Souza."


# ── extrair_texto ──────────────────────────────────────────────────────────


def test_extrair_texto_de_paragrafo_simples():
    docx = _docx_simples("Reu: Janaina Pereira da Silva Matos.")
    assert extrair_texto(docx) == "Reu: Janaina Pereira da Silva Matos."


def test_extrair_texto_de_multiplos_paragrafos_preserva_ordem():
    doc = Document()
    doc.add_paragraph("Primeiro paragrafo.")
    doc.add_paragraph("Segundo paragrafo.")
    doc.add_paragraph("Terceiro paragrafo.")
    out = BytesIO()
    doc.save(out)

    texto = extrair_texto(out.getvalue())
    assert texto == "Primeiro paragrafo.\nSegundo paragrafo.\nTerceiro paragrafo."


def test_extrair_texto_inclui_celulas_de_tabela():
    docx = _docx_com_tabela(
        [
            ["Campo", "Valor"],
            ["Reclamado", "Janaina Pereira"],
        ]
    )
    texto = extrair_texto(docx)

    assert "Campo" in texto
    assert "Valor" in texto
    assert "Reclamado" in texto
    assert "Janaina Pereira" in texto


def test_extrair_texto_ignora_paragrafos_vazios():
    doc = Document()
    doc.add_paragraph("Um.")
    doc.add_paragraph("")
    doc.add_paragraph("Tres.")
    out = BytesIO()
    doc.save(out)

    assert extrair_texto(out.getvalue()) == "Um.\nTres."


def test_extrair_texto_bytes_vazios_levanta_erro():
    with pytest.raises(SubstituicaoError):
        extrair_texto(b"")


def test_extrair_texto_bytes_invalidos_levantam_erro():
    with pytest.raises(SubstituicaoError):
        extrair_texto(b"nao eh um docx valido")
