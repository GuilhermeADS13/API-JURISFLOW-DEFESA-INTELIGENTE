"""Testes do PR6 #1 — Long Context (pre-filtragem + truncamento inteligente).

Cobre:
- Texto < limite passa direto sem mexer.
- Pre-filtragem reconhece secoes juridicas comuns (DOS FATOS, DO DIREITO, DOS PEDIDOS).
- Truncamento inteligente preserva inicio + fim quando pre-filtragem nao basta.
- Anexos consolidados respeitam o limite global.
"""

from __future__ import annotations

import pytest

from App.services.peticao_extractor import (
    MAX_TEXTO_PETICAO_CHARS,
    _aplicar_long_context,
    prefiltrar_secoes_juridicas,
)


PETICAO_CURTA = """
PETICAO INICIAL

I - DOS FATOS

O Reclamante Joao da Silva trabalhou de 2020 a 2024 sem registro em CTPS.
Cumpriu jornada de 8 horas/dia.

II - DO DIREITO

Aplica-se o art. 7, XV, da CF/88, garantindo o repouso semanal remunerado.
A CLT, art. 442, define vinculo empregaticio.

III - DOS PEDIDOS

a) Reconhecimento do vinculo de emprego;
b) Pagamento de horas extras nao quitadas;
c) Adicional noturno pelo periodo trabalhado.

Valor da causa: R$ 27.598,41.
"""


def test_prefiltrar_reconhece_secoes_principais():
    secoes = prefiltrar_secoes_juridicas(PETICAO_CURTA)
    assert "FATOS" in secoes
    assert "DIREITO" in secoes
    assert "PEDIDOS" in secoes
    # Conteudo da secao FATOS captura ate o cabecalho seguinte.
    assert "Joao da Silva" in secoes["FATOS"]
    assert "DO DIREITO" not in secoes["FATOS"]  # nao vaza para a proxima


def test_prefiltrar_aceita_numeracao_arabe():
    texto = """
1. DOS FATOS
fato A.
2. DO DIREITO
direito B.
3. DOS PEDIDOS
pedido C.
"""
    secoes = prefiltrar_secoes_juridicas(texto)
    assert "FATOS" in secoes and "fato A" in secoes["FATOS"]
    assert "DIREITO" in secoes and "direito B" in secoes["DIREITO"]
    assert "PEDIDOS" in secoes and "pedido C" in secoes["PEDIDOS"]


def test_prefiltrar_texto_sem_secoes_retorna_dict_vazio():
    texto = "Texto livre sem cabecalhos juridicos comuns. Apenas paragrafos soltos."
    secoes = prefiltrar_secoes_juridicas(texto)
    assert secoes == {}


def test_prefiltrar_concatena_secoes_repetidas():
    texto = """
DOS PEDIDOS
pedido 1
DOS FATOS
fato A
DOS PEDIDOS
pedido 2
"""
    secoes = prefiltrar_secoes_juridicas(texto)
    # PEDIDOS aparece 2x - ambos capturados
    assert "pedido 1" in secoes["PEDIDOS"]
    assert "pedido 2" in secoes["PEDIDOS"]


def test_aplicar_long_context_texto_pequeno_devolve_inalterado_se_passar_de_limite():
    """Quando o texto e menor que o limite, _aplicar_long_context nao deve ser chamado;
    mas se for chamado mesmo assim, retorna conteudo util sem quebrar."""
    # Texto pequeno simulando passagem (limite=100 bem baixo).
    texto = PETICAO_CURTA
    out = _aplicar_long_context(texto, limite=10_000)
    # Cabe no limite com folga, deve retornar com pre-filtro aplicado (com cabecalho).
    assert "PEDIDOS" in out
    assert "FATOS" in out
    assert len(out) <= 10_000


def test_aplicar_long_context_prefiltragem_quando_excede_limite():
    """Texto enorme com secoes claras: pre-filtragem deve manter as secoes priorizadas."""
    preambulo = "preambulo irrelevante " * 5000  # ~100k chars de ruido
    pedidos = "pedido critico que NAO pode sumir do contexto. " * 10
    texto = (
        preambulo
        + "\n\nI - DOS FATOS\nfatos importantes.\n\n"
        + "II - DO DIREITO\nfundamentos juridicos.\n\n"
        + "III - DOS PEDIDOS\n"
        + pedidos
    )

    out = _aplicar_long_context(texto, limite=20_000)

    assert len(out) <= 20_000
    # Pedidos NAO podem ter sumido — sao a secao mais critica.
    assert "pedido critico que NAO pode sumir" in out


def test_aplicar_long_context_truncamento_quando_secoes_nao_ajudam():
    """Texto enorme sem cabecalhos: cai para truncamento inicio+fim com elipse."""
    texto = "frase de meio " * 20_000  # 280k chars sem secoes
    inicio_marker = "INICIO_DO_DOCUMENTO"
    fim_marker = "FIM_DO_DOCUMENTO_PEDIDOS_AQUI"
    texto_completo = inicio_marker + " " + texto + " " + fim_marker

    out = _aplicar_long_context(texto_completo, limite=10_000)

    assert len(out) <= 10_000
    # Inicio e fim preservados.
    assert inicio_marker in out
    assert fim_marker in out
    # Elipse central explicita.
    assert "[...trecho intermediario omitido" in out


def test_extrair_texto_peticao_curta_passa_direto():
    """Texto pequeno (<<MAX) nao deve ser pre-filtrado."""
    from io import BytesIO

    from docx import Document

    doc = Document()
    for linha in PETICAO_CURTA.splitlines():
        doc.add_paragraph(linha)
    out = BytesIO()
    doc.save(out)

    from App.services.peticao_extractor import extrair_texto_peticao
    texto = extrair_texto_peticao(out.getvalue(), "peticao.docx")
    # Texto pequeno volta integralmente sem cabecalho de pre-filtragem.
    assert "Joao da Silva" in texto
    assert "[...preambulo omitido" not in texto


def test_extrair_texto_peticao_grande_aplica_long_context():
    """Texto > MAX deve passar por pre-filtragem ou truncamento."""
    from io import BytesIO

    from docx import Document

    doc = Document()
    # Cabecalho identificavel
    doc.add_paragraph("Processo 0000123-45.2026.5.06.0341 — Reclamante Joao")
    # Preambulo gigante sem cabecalho de secao
    for _ in range(800):
        doc.add_paragraph("paragrafo de preambulo irrelevante para a contestacao.")
    # Secoes priorizadas no fim
    doc.add_paragraph("I - DOS FATOS")
    doc.add_paragraph("Fato critico que precisa chegar ao Claude.")
    doc.add_paragraph("II - DOS PEDIDOS")
    doc.add_paragraph("Pedido critico que NAO pode ser truncado.")
    out = BytesIO()
    doc.save(out)

    from App.services.peticao_extractor import extrair_texto_peticao
    texto = extrair_texto_peticao(out.getvalue(), "peticao.docx")

    assert len(texto) <= MAX_TEXTO_PETICAO_CHARS
    # Pedidos preservados via pre-filtragem.
    assert "Pedido critico que NAO pode ser truncado" in texto
    # Cabecalho do processo preservado no inicio.
    assert "0000123-45.2026.5.06.0341" in texto
