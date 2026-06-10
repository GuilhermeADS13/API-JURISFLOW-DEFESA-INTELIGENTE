"""Testes do contestacao_docx_builder, com foco no ROL DE DOCUMENTOS (PR14).

Strategy: invocar `montar_docx_programatico` com diferentes shapes de
`documentos_anexos` e inspecionar os paragrafos do .docx resultante.
"""

from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document


def _dados_minimos():
    return {
        "numero_processo": "0001234-56.2026.5.06.0001",
        "autor": "Joao da Silva",
        "reu": "Empresa XYZ LTDA",
        "tipo_acao": "Trabalhista",
        "vara": "1a Vara do Trabalho de Recife",
    }


def _minuta_minima(**extra):
    base = {
        "tese_central": "Improcedencia",
        "preliminares": "Prescricao quinquenal.",
        "merito": "Os pedidos carecem de fundamento.",
        "fundamentos": "Art. 818 CLT.",
        "pedidos": "Improcedencia total.",
    }
    base.update(extra)
    return base


def _texto_do_docx(bytes_docx: bytes) -> str:
    """Concatena todos os paragrafos do docx pra busca por substring."""
    doc = Document(BytesIO(bytes_docx))
    return "\n".join(p.text for p in doc.paragraphs)


# ─────────────────────────────────────────────────────────────────────────────
# PR14 — Rol de Documentos
# ─────────────────────────────────────────────────────────────────────────────


def test_rol_de_documentos_aparece_quando_anexos_preenchidos():
    from App.services.contestacao_docx_builder import montar_docx_programatico

    anexos = [
        {
            "numero": "Doc. 01",
            "tipo": "Folha de Ponto",
            "descricao": "Cartoes de ponto do periodo de 01/2020 a 11/2025.",
        },
        {
            "numero": "Doc. 02",
            "tipo": "Extrato FGTS",
            "descricao": "Extrato analitico da CEF do periodo contratual.",
        },
    ]

    docx_bytes = montar_docx_programatico(
        _dados_minimos(),
        _minuta_minima(documentos_anexos=anexos),
    )
    texto = _texto_do_docx(docx_bytes)

    assert "ROL DE DOCUMENTOS" in texto
    assert "Folha de Ponto" in texto
    assert "Extrato FGTS" in texto
    # Cada item gera placeholder visual pro advogado
    assert texto.count("[ANEXAR ARQUIVO]") == 2
    # Descricoes presentes
    assert "Cartoes de ponto do periodo" in texto
    assert "Extrato analitico" in texto


def test_rol_omitido_quando_anexos_ausente():
    from App.services.contestacao_docx_builder import montar_docx_programatico

    # Sem campo documentos_anexos
    docx_bytes = montar_docx_programatico(_dados_minimos(), _minuta_minima())
    texto = _texto_do_docx(docx_bytes)

    assert "ROL DE DOCUMENTOS" not in texto
    assert "[ANEXAR ARQUIVO]" not in texto


def test_rol_omitido_quando_lista_vazia():
    from App.services.contestacao_docx_builder import montar_docx_programatico

    docx_bytes = montar_docx_programatico(
        _dados_minimos(), _minuta_minima(documentos_anexos=[])
    )
    texto = _texto_do_docx(docx_bytes)

    assert "ROL DE DOCUMENTOS" not in texto


def test_rol_descarta_itens_malformados():
    from App.services.contestacao_docx_builder import montar_docx_programatico

    anexos = [
        {"numero": "Doc. 01", "tipo": "Folha de Ponto", "descricao": "Valida."},
        "string solta — invalida",
        {},  # dict sem tipo nem descricao — invalido
        None,  # invalido
        {"tipo": "TRCT", "descricao": "valida sem numero"},  # valido (numero default)
    ]

    docx_bytes = montar_docx_programatico(
        _dados_minimos(), _minuta_minima(documentos_anexos=anexos)
    )
    texto = _texto_do_docx(docx_bytes)

    # 2 itens validos -> 2 placeholders
    assert texto.count("[ANEXAR ARQUIVO]") == 2
    assert "Folha de Ponto" in texto
    assert "TRCT" in texto


def test_rol_cap_em_10_itens():
    from App.services.contestacao_docx_builder import montar_docx_programatico

    anexos = [
        {"numero": f"Doc. {i:02d}", "tipo": f"Tipo {i}", "descricao": f"Desc {i}."}
        for i in range(1, 16)  # 15 itens — alem do cap
    ]

    docx_bytes = montar_docx_programatico(
        _dados_minimos(), _minuta_minima(documentos_anexos=anexos)
    )
    texto = _texto_do_docx(docx_bytes)

    # Cap em 10 itens — 11+ nao aparecem
    assert texto.count("[ANEXAR ARQUIVO]") == 10
    assert "Tipo 10" in texto
    assert "Tipo 11" not in texto


def test_rol_aceita_so_descricao_sem_tipo():
    from App.services.contestacao_docx_builder import montar_docx_programatico

    anexos = [
        {"descricao": "Documento generico sem tipo formal."},
    ]

    docx_bytes = montar_docx_programatico(
        _dados_minimos(), _minuta_minima(documentos_anexos=anexos)
    )
    texto = _texto_do_docx(docx_bytes)

    assert "Documento generico sem tipo formal." in texto
    assert texto.count("[ANEXAR ARQUIVO]") == 1


def test_rol_numera_default_quando_numero_ausente():
    from App.services.contestacao_docx_builder import montar_docx_programatico

    anexos = [
        {"tipo": "A", "descricao": "primeiro"},
        {"tipo": "B", "descricao": "segundo"},
    ]

    docx_bytes = montar_docx_programatico(
        _dados_minimos(), _minuta_minima(documentos_anexos=anexos)
    )
    texto = _texto_do_docx(docx_bytes)

    # numeracao default 'Doc. 01', 'Doc. 02'
    assert "Doc. 01" in texto
    assert "Doc. 02" in texto


def test_rol_nao_renderiza_quando_anexos_nao_eh_lista():
    """Robustez: se Claude retornar dict ou string em vez de list, NAO quebra."""
    from App.services.contestacao_docx_builder import montar_docx_programatico

    for valor_invalido in ("string", {"foo": "bar"}, 42):
        docx_bytes = montar_docx_programatico(
            _dados_minimos(),
            _minuta_minima(documentos_anexos=valor_invalido),
        )
        texto = _texto_do_docx(docx_bytes)
        assert "ROL DE DOCUMENTOS" not in texto


# ─────────────────────────────────────────────────────────────────────────────
# PR15 — Embedding de imagens no ROL
# ─────────────────────────────────────────────────────────────────────────────


def _png_colorido(cor: str = "white") -> bytes:
    """PNG 2x2 com cor especifica. Cores diferentes evitam dedup do python-docx."""
    from PIL import Image

    buf = BytesIO()
    Image.new("RGB", (2, 2), color=cor).save(buf, format="PNG")
    return buf.getvalue()


# Contador global pra gerar PNGs unicos em cada chamada (cor sequencial).
_CORES_TESTE = ("red", "green", "blue", "yellow", "purple", "orange", "cyan", "magenta", "pink", "brown")


def _imagem_embedavel(tipo: str, nome: str = "fake.png", cor_idx: int | None = None):
    """Helper pra criar ImagemEmbedavel com PNG UNICO (evita dedup do docx)."""
    from App.services.embed_processor import ImagemEmbedavel

    # Cor diferente por chamada — se cor_idx for None, hash do nome decide.
    if cor_idx is None:
        cor_idx = hash(nome) % len(_CORES_TESTE)
    cor = _CORES_TESTE[cor_idx % len(_CORES_TESTE)]
    return ImagemEmbedavel(
        tipo=tipo,
        nome=nome,
        bytes_png=_png_colorido(cor),
        pagina=1,
        eh_imagem_direta=True,
    )


def _contar_imagens_no_docx(bytes_docx: bytes) -> int:
    """Conta arquivos em word/media/ — onde python-docx grava imagens embedded."""
    import zipfile

    with zipfile.ZipFile(BytesIO(bytes_docx)) as z:
        return sum(1 for name in z.namelist() if name.startswith("word/media/"))


def test_imagem_embeddada_quando_tipo_casa_com_anexo():
    from App.services.contestacao_docx_builder import montar_docx_programatico

    anexos = [
        {"numero": "Doc. 01", "tipo": "Folha de Ponto", "descricao": "jornada"},
        {"numero": "Doc. 02", "tipo": "Extrato FGTS", "descricao": "depositos"},
    ]
    imagens = [
        _imagem_embedavel("folha_ponto", "ponto.png"),
        _imagem_embedavel("fgts", "fgts.png"),
    ]

    docx_bytes = montar_docx_programatico(
        _dados_minimos(),
        _minuta_minima(documentos_anexos=anexos),
        imagens_embedar=imagens,
    )

    # Ambas imagens embedded — nenhum placeholder fallback
    assert _contar_imagens_no_docx(docx_bytes) == 2
    texto = _texto_do_docx(docx_bytes)
    assert "[ANEXAR ARQUIVO]" not in texto
    # Legendas com nome do arquivo aparecem
    assert "ponto.png" in texto
    assert "fgts.png" in texto


def test_placeholder_quando_tipo_nao_casa():
    """Anexo declarado pelo Claude mas sem imagem correspondente: placeholder."""
    from App.services.contestacao_docx_builder import montar_docx_programatico

    anexos = [
        {"numero": "Doc. 01", "tipo": "Folha de Ponto", "descricao": "jornada"},
        {"numero": "Doc. 02", "tipo": "Extrato FGTS", "descricao": "depositos"},
    ]
    # So uma imagem (folha_ponto) — FGTS vai cair em placeholder
    imagens = [_imagem_embedavel("folha_ponto", "ponto.png")]

    docx_bytes = montar_docx_programatico(
        _dados_minimos(),
        _minuta_minima(documentos_anexos=anexos),
        imagens_embedar=imagens,
    )

    assert _contar_imagens_no_docx(docx_bytes) == 1
    texto = _texto_do_docx(docx_bytes)
    # 1 placeholder ainda (pro FGTS sem imagem)
    assert texto.count("[ANEXAR ARQUIVO]") == 1


def test_imagens_nao_casadas_viram_outras_provas():
    """Imagens enviadas sem item correspondente no ROL caem em bloco extra."""
    from App.services.contestacao_docx_builder import montar_docx_programatico

    anexos = [
        {"numero": "Doc. 01", "tipo": "Contrato de Trabalho", "descricao": "x"},
    ]
    # Imagem de tipo que nao casa com nenhum item declarado
    imagens = [
        _imagem_embedavel("laudo_pericial", "laudo.png"),
        _imagem_embedavel("print", "email.png"),
    ]

    docx_bytes = montar_docx_programatico(
        _dados_minimos(),
        _minuta_minima(documentos_anexos=anexos),
        imagens_embedar=imagens,
    )

    # Total: 0 casadas no ROL + 2 em OUTRAS PROVAS = 2 imagens embedded
    assert _contar_imagens_no_docx(docx_bytes) == 2
    texto = _texto_do_docx(docx_bytes)
    assert "OUTRAS PROVAS ANEXAS" in texto


def test_normalizar_tipo_anexo_casa_variantes_humanas():
    from App.services.contestacao_docx_builder import _normalizar_tipo_anexo

    assert _normalizar_tipo_anexo("Folha de Ponto") == "folha_ponto"
    assert _normalizar_tipo_anexo("Cartoes de Ponto") == "folha_ponto"
    assert _normalizar_tipo_anexo("Extrato FGTS") == "fgts"
    assert _normalizar_tipo_anexo("Extrato Analitico FGTS") == "fgts"
    assert _normalizar_tipo_anexo("TRCT") == "trct"
    assert _normalizar_tipo_anexo("Termo de Rescisao") == "trct"
    assert _normalizar_tipo_anexo("Laudo Pericial") == "laudo_pericial"
    assert _normalizar_tipo_anexo("PPP") == "laudo_pericial"
    assert _normalizar_tipo_anexo("Contrato de Trabalho") == "contrato"
    assert _normalizar_tipo_anexo("CTPS Digital") == "ctps"
    assert _normalizar_tipo_anexo("Prints de e-mail") == "print"
    # Fallback
    assert _normalizar_tipo_anexo("Coisa Aleatoria") == "outro"
    assert _normalizar_tipo_anexo("") == "outro"
    assert _normalizar_tipo_anexo(None) == "outro"


def test_falha_em_embedding_de_imagem_corrompida_nao_quebra():
    """Imagem com bytes invalidos vira placeholder de erro, nao crasha o docx."""
    from App.services.contestacao_docx_builder import montar_docx_programatico
    from App.services.embed_processor import ImagemEmbedavel

    imagem_quebrada = ImagemEmbedavel(
        tipo="fgts",
        nome="quebrada.png",
        bytes_png=b"nao_eh_png_valido",
        pagina=1,
        eh_imagem_direta=True,
    )

    docx_bytes = montar_docx_programatico(
        _dados_minimos(),
        _minuta_minima(documentos_anexos=[{"numero": "Doc. 01", "tipo": "FGTS", "descricao": "x"}]),
        imagens_embedar=[imagem_quebrada],
    )
    texto = _texto_do_docx(docx_bytes)
    # Nenhuma imagem embedded mas docx gerou OK + placeholder de erro
    assert _contar_imagens_no_docx(docx_bytes) == 0
    assert "FALHA AO EMBEDAR" in texto


def test_rol_so_renderiza_se_houver_anexos_ou_imagens():
    """Sem anexos no JSON E sem imagens, secao NAO aparece."""
    from App.services.contestacao_docx_builder import montar_docx_programatico

    docx_bytes = montar_docx_programatico(_dados_minimos(), _minuta_minima())
    texto = _texto_do_docx(docx_bytes)
    assert "ROL DE DOCUMENTOS" not in texto
    assert "OUTRAS PROVAS" not in texto


def test_sem_anexos_mas_com_imagens_renderiza_outras_provas():
    """Imagens sem ROL: bloco OUTRAS PROVAS aparece direto."""
    from App.services.contestacao_docx_builder import montar_docx_programatico

    imagens = [_imagem_embedavel("outro", "anexo.png")]
    docx_bytes = montar_docx_programatico(
        _dados_minimos(), _minuta_minima(), imagens_embedar=imagens
    )
    texto = _texto_do_docx(docx_bytes)
    assert "OUTRAS PROVAS ANEXAS" in texto
    assert _contar_imagens_no_docx(docx_bytes) == 1


# ─────────────────────────────────────────────────────────────────────────────
# PR16.2 — regressao: timbre/header/footer preservados quando ha documentos_anexos
# ─────────────────────────────────────────────────────────────────────────────


def _modelo_b64_sem_heading_styles() -> str:
    """Cria um .docx modelo MINIMO com header/footer customizado mas SEM
    os styles 'Heading 1/2/3' definidos. Reproduz o modelo do escritorio
    G. Trindade que causou a regressao da peca #44.
    """
    import base64

    from docx import Document
    from docx.shared import Pt

    doc = Document()
    # Adiciona header com "TIMBRE DO ESCRITORIO"
    section = doc.sections[0]
    h = section.header.paragraphs[0]
    h.text = "TIMBRE G. TRINDADE ADVOGADOS"
    section.footer.paragraphs[0].text = "Rodape do escritorio"
    # 1 paragrafo no body (sera limpo pelo template builder)
    doc.add_paragraph("placeholder")
    # REMOVE os styles 'Heading X' do template — simula modelo do escritorio
    # antigo que so tem o style 'Normal'
    styles = doc.styles
    for nome in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        try:
            styles[nome].delete()
        except (KeyError, Exception):
            pass

    out = BytesIO()
    doc.save(out)
    return base64.b64encode(out.getvalue()).decode("ascii")


def test_pr162_rol_documentos_nao_quebra_template_sem_heading_styles():
    """REGRESSAO da peca #44: builder template caia em fallback porque
    doc.add_heading('...', level=2) levantava KeyError quando o modelo
    do escritorio nao tinha o style 'Heading 2'.
    """
    from App.services.contestacao_docx_builder import montar_docx_com_modelo

    minuta = _minuta_minima(documentos_anexos=[
        {"numero": "Doc. 01", "tipo": "Folha de Ponto", "descricao": "jornada"},
    ])

    modelo_b64 = _modelo_b64_sem_heading_styles()
    docx_bytes = montar_docx_com_modelo(modelo_b64, _dados_minimos(), minuta)

    # Pre-fix do PR16.2: retornaria None (KeyError em _safe_step) e cairia
    # no fallback programatico — perdendo o timbre.
    assert docx_bytes is not None, (
        "Builder template retornou None — provavelmente caiu em fallback "
        "por KeyError do Heading 2 (regressao PR14/15 pre-fix PR16.2)"
    )
    # Header preservado (timbre do escritorio aparece no docx final)
    from docx import Document
    d = Document(BytesIO(docx_bytes))
    header_text = "\n".join(p.text for p in d.sections[0].header.paragraphs)
    assert "TIMBRE G. TRINDADE" in header_text, (
        "Header do modelo (timbre) sumiu — _build_docx_from_template provavelmente "
        "limpou tudo ou caiu no fallback programatico"
    )
    # ROL DE DOCUMENTOS renderizado mesmo sem o style 'Heading 2'
    body_text = "\n".join(p.text for p in d.paragraphs)
    assert "ROL DE DOCUMENTOS" in body_text


def test_pr162_outras_provas_nao_quebra_template_sem_heading_styles():
    """Mesma regressao com bloco OUTRAS PROVAS ANEXAS (sem documentos_anexos mas
    com imagens_embedar que nao casam)."""
    from App.services.contestacao_docx_builder import montar_docx_com_modelo

    imagens = [_imagem_embedavel("outro", "anexo.png")]
    modelo_b64 = _modelo_b64_sem_heading_styles()
    docx_bytes = montar_docx_com_modelo(
        modelo_b64,
        _dados_minimos(),
        _minuta_minima(),
        imagens_embedar=imagens,
    )

    assert docx_bytes is not None, "Builder template nao deve cair em fallback"
    from docx import Document
    d = Document(BytesIO(docx_bytes))
    header_text = "\n".join(p.text for p in d.sections[0].header.paragraphs)
    assert "TIMBRE G. TRINDADE" in header_text


# ─────────────────────────────────────────────────────────────────────────────
# PR16.3 — line_spacing cap subido p/ 1.25 (fidelidade ao modelo G. Trindade)
# ─────────────────────────────────────────────────────────────────────────────


def test_pr163_cap_line_spacing_preserva_125_do_modelo():
    """cap_line_spacing(1.25) deve devolver 1.25, nao 1.20 (cap antigo do PR
    de 02/06). Modelos do escritorio que usam 1.25 nao podem ser empurrados
    pra 1.20 — isso era a causa de o docx gerado renderizar diferente.
    """
    from App.services.docx_style_defaults import (
        LINE_SPACING_CAP_FROM_TEMPLATE,
        cap_line_spacing,
    )

    assert LINE_SPACING_CAP_FROM_TEMPLATE == 1.25
    assert cap_line_spacing(1.25) == 1.25
    assert cap_line_spacing(1.20) == 1.20
    assert cap_line_spacing(1.15) == 1.15
    # Modelos com 1.5 ou 2.0 continuam capeados em 1.25 (defesa contra
    # templates legados Word com line_spacing exagerado)
    assert cap_line_spacing(1.50) == 1.25
    assert cap_line_spacing(2.00) == 1.25
