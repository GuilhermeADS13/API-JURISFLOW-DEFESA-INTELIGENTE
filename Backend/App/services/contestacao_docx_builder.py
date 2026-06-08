"""Constroi o .docx final da contestacao a partir da minuta JSON do agente.

Dois caminhos:
- `montar_docx_com_modelo`: abre o modelo base do escritorio com python-docx,
  limpa o body preservando headers/footers/watermark/sectPr e insere paragrafos
  novos com a tipografia do escritorio (Arial 11 justificado, numeracao 01.-,
  headings negrito+sublinhado, Markdown inline **bold**/__underline__/quote).
- `montar_docx_programatico`: gera um .docx do zero com python-docx quando nao
  ha modelo base, estruturado em secoes (CONTESTACAO, TESE, PRELIMINARES,
  MERITO, IMPUGNACAO, FUNDAMENTOS, PEDIDOS).

Historico: na Etapa 5 separaram-se helpers para reduzir CC. No PR de download
de peca (commit b78a2a0) o caminho com modelo migrou de docxtpl Subdoc (que
duplicava conteudo) para python-docx puro com renderer de Markdown proprio.
"""

from __future__ import annotations

import base64
import binascii
import logging
import re
from contextvars import ContextVar
from copy import deepcopy
from datetime import datetime
from io import BytesIO
from types import MappingProxyType
from typing import Any, Callable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from App.services.docx_style_defaults import (
    FONT_NAME_DEFAULT,
    FONT_SIZE_PT_DEFAULT,
    LINE_SPACING_DEFAULT,
    SPACE_AFTER_PT_DEFAULT,
    SPACE_BEFORE_SECAO1_PT_DEFAULT,
    SPACE_BEFORE_SECAO2_PT_DEFAULT,
    aplicar_espacamento_padrao as _aplicar_espacamento_padrao_central,
    cap_font_size_pt,
    cap_line_spacing,
    cap_space_after_pt,
)

logger = logging.getLogger(__name__)


# ──────────────────────── Estilo dinamico do template ─────────────────────────
# Em vez de hardcodar fonte/espaco/tab para o modelo G. Trindade, extraimos
# essas propriedades do modelo base que o usuario subiu. Os helpers leem essas
# configs via ContextVar (thread-safe para requests concorrentes).

_ESTILO_PADRAO: dict[str, Any] = {
    # IMPORTANTE: valores padroes de tipografia/espacamento sao a unica fonte
    # da verdade em `App/services/docx_style_defaults.py`. Qualquer ajuste
    # vai LA, nao aqui — assim outros builders (peticao, parecer, etc) ficam
    # consistentes sem repetir o bug do PDF LibreOffice fora de paginacao.
    "font_name": FONT_NAME_DEFAULT,
    "font_size_pt": FONT_SIZE_PT_DEFAULT,
    "line_spacing": LINE_SPACING_DEFAULT,
    "space_after_pt": SPACE_AFTER_PT_DEFAULT,
    "space_before_secao1_pt": SPACE_BEFORE_SECAO1_PT_DEFAULT,
    "space_before_secao2_pt": SPACE_BEFORE_SECAO2_PT_DEFAULT,
    # Tab stop default do Word eh ~1.27cm. Modelo G. Trindade nao define
    # tab_stops customizados — usa 'NN.- TAB ESPACO TAB' (gera ~2-3cm).
    "tab_stop_cm": 1.27,
    # Recuo do blockquote (citacoes de lei/sumula): 3cm igual ao modelo.
    "quote_left_indent_cm": 3.0,
}

_ESTILO_ATIVO: ContextVar[dict[str, Any]] = ContextVar(
    "_ESTILO_ATIVO", default=_ESTILO_PADRAO
)


def _estilo() -> Any:
    """Retorna view read-only do estilo ativo (do modelo base) ou do default.

    Usa MappingProxyType pra impedir mutacao acidental que poluiria o
    _ESTILO_PADRAO module-level entre requests.
    """
    valor = _ESTILO_ATIVO.get()
    if valor is None:
        valor = _ESTILO_PADRAO
    return MappingProxyType(valor)


def _extrair_estilo_modelo(doc: Any) -> dict[str, Any]:
    """Inspeciona o template e extrai fonte/espacamento dominantes.

    Estrategia:
    1. Le style 'Normal' (fallback do template).
    2. Sobrescreve com o primeiro paragrafo do body que tenha texto.
    3. Em caso de ausencia, retorna defaults conservadores (Arial 11, 1.25, 6pt).
    """
    estilo = dict(_ESTILO_PADRAO)
    try:
        normal = doc.styles["Normal"]
        if getattr(normal.font, "name", None):
            estilo["font_name"] = normal.font.name
        if getattr(normal.font, "size", None):
            estilo["font_size_pt"] = cap_font_size_pt(normal.font.size.pt)
    except Exception as err:  # noqa: BLE001 — template estranho, segue com default
        logger.debug("Sem acesso ao style Normal do template: %s", err)

    for paragrafo in doc.paragraphs:
        if not paragrafo.text.strip():
            continue
        # Cada paragrafo eh envelopado em try/except: template degenerado
        # com runs/paragraph_format invalidos nao deve quebrar a geracao.
        try:
            for run in paragrafo.runs:
                if getattr(run.font, "name", None):
                    estilo["font_name"] = run.font.name
                    break
            for run in paragrafo.runs:
                if getattr(run.font, "size", None):
                    # Cap em [11, 12]pt — protege contra templates com
                    # tamanho exotico (ver docx_style_defaults).
                    estilo["font_size_pt"] = cap_font_size_pt(run.font.size.pt)
                    break
            pf = paragrafo.paragraph_format
            ls = pf.line_spacing
            # ls pode ser float (multiplicador) ou WD_LINE_SPACING enum (int).
            # So aceita valores numericos positivos plausiveis (0.5–3.0).
            # Caps em `docx_style_defaults` evitam que templates antigos
            # do escritorio (1.5 = padrao Word legado) gerem PDF com
            # paginacao diferente entre Word e LibreOffice.
            if isinstance(ls, (int, float)) and 0.5 <= float(ls) <= 3.0:
                estilo["line_spacing"] = cap_line_spacing(ls)
            if pf.space_after:
                try:
                    estilo["space_after_pt"] = cap_space_after_pt(
                        pf.space_after.pt
                    )
                except (AttributeError, TypeError):
                    pass
        except Exception as err:  # noqa: BLE001 — paragrafo corrompido, ignora
            logger.debug("Paragrafo do template ignorado: %s", err)
            continue
        # Maior tab stop = TAB grande do "01.- texto" do escritorio
        maior_tab = 0.0
        for tab in pf.tab_stops:
            if tab.position:
                try:
                    cm = float(tab.position.cm)
                    if cm > maior_tab:
                        maior_tab = cm
                except (AttributeError, TypeError):
                    continue
        if maior_tab > 0:
            estilo["tab_stop_cm"] = maior_tab
        break

    # Procura recuo de blockquote em paragrafos com left_indent (citacoes de lei)
    for paragrafo in doc.paragraphs:
        pf = paragrafo.paragraph_format
        if pf.left_indent:
            try:
                cm = float(pf.left_indent.cm)
                if 1.0 < cm < 6.0:  # filtro: recuo razoavel
                    estilo["quote_left_indent_cm"] = cm
                    break
            except (AttributeError, TypeError):
                continue

    logger.info(
        "Estilo extraido do template: font=%s %.1fpt, line=%.2f, after=%.1fpt, tab=%.1fcm",
        estilo["font_name"],
        estilo["font_size_pt"],
        estilo["line_spacing"],
        estilo["space_after_pt"],
        estilo["tab_stop_cm"],
    )
    return estilo


# ─────────────────────── Geracao programatica do .docx ───────────────────────


def montar_docx_programatico(
    dados: dict[str, Any],
    minuta: dict[str, Any],
    *,
    imagens_embedar: list | None = None,
) -> bytes:
    """Gera .docx do zero quando nao ha modelo base do escritorio.

    Estrutura espelha a do Node 6 do guia tecnico v2: cabecalho com partes,
    tese central, preliminares, merito, impugnacao por pedido, fundamentos,
    pedidos. Tudo em portugues.

    PR15 — `imagens_embedar` (list[ImagemEmbedavel] do embed_processor) sao
    inseridas como imagens reais no ROL DE DOCUMENTOS final quando o tipo
    casa com algum item de minuta['documentos_anexos'].
    """
    doc = Document()
    _aplicar_estilo_base(doc)
    _escrever_cabecalho(doc, dados)
    _escrever_secoes_minuta(doc, minuta, imagens_embedar=imagens_embedar)
    _escrever_rodape(doc, minuta)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _aplicar_estilo_base(doc: Any) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)


def _escrever_cabecalho(doc: Any, dados: dict[str, Any]) -> None:
    """Titulo + identificacao das partes."""
    titulo = doc.add_heading("CONTESTACAO", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_paragraph(doc, f"Processo numero: {dados.get('numero_processo') or 'a definir'}")
    _add_paragraph(doc, f"Autor: {dados.get('autor') or ''}")
    _add_paragraph(doc, f"Reu: {dados.get('reu') or ''}")
    _add_paragraph(doc, f"Tipo de acao: {dados.get('tipo_acao') or ''}")
    if dados.get("vara"):
        _add_paragraph(doc, f"Vara: {dados.get('vara')}")


# Tabela das secoes da minuta no formato (chave_minuta, titulo_no_docx).
# Usada pelo loop de _escrever_secoes_minuta — adicionar uma secao nova so
# requer estender essa lista (e o renderer especial para impugnacao, se for
# por pedido). Mantem o codigo aberto a extensao, fechado a modificacao.
_SECOES_TEXTO = (
    ("tese_central", "I — TESE CENTRAL"),
    ("preliminares", "II — PRELIMINARES"),
    ("merito", "III — DO MERITO"),
)
_SECOES_FINAIS = (
    ("fundamentos", "IV — FUNDAMENTOS JURIDICOS"),
    ("pedidos", "V — PEDIDOS"),
)


def _escrever_secoes_minuta(
    doc: Any,
    minuta: dict[str, Any],
    *,
    imagens_embedar: list | None = None,
) -> None:
    """Itera secoes da minuta na ordem definida em _SECOES_TEXTO.

    impugnacao_pedidos NAO eh renderizada como secao romana separada:
    o merito ja deve conter as impugnacoes. Se vier dict nao-vazio,
    aparece como subsecao do merito (III.Z).

    Apos PEDIDOS (V), renderiza o ROL DE DOCUMENTOS (VI) quando o Claude
    populou `documentos_anexos[]`. Se houver `imagens_embedar` (PR15), as
    imagens reais entram no lugar do placeholder [ANEXAR ARQUIVO].
    """
    for chave, titulo in _SECOES_TEXTO:
        _escrever_secao_texto(doc, titulo, minuta.get(chave))

    # impugnacao_pedidos eh referencia interna do prompt — nao vai na peca.

    for chave, titulo in _SECOES_FINAIS:
        _escrever_secao_texto(doc, titulo, minuta.get(chave))

    # PR14/PR15 — Rol de documentos probatorios (com imagens embedded se houver).
    _escrever_rol_documentos(
        doc, minuta.get("documentos_anexos"), imagens_embedar=imagens_embedar
    )


def _normalizar_tipo_anexo(tipo_bruto: Any) -> str:
    """Normaliza tipo declarado pelo Claude em chave canonica PR15.

    Aceita varios formatos: 'Folha de Ponto' -> 'folha_ponto'; 'FGTS' -> 'fgts';
    'Laudo Pericial Tecnico' -> 'laudo_pericial' (matching por substring).
    Retorna 'outro' se nao casar nenhum padrao.
    """
    t = str(tipo_bruto or "").lower().strip()
    if not t:
        return "outro"
    # Match por substring — robusto ao Claude usar variacoes ("Folha de Ponto" vs "Cartoes de Ponto")
    if any(kw in t for kw in ("folha de ponto", "cartao", "cartoes", "controle de jornada")):
        return "folha_ponto"
    if "fgts" in t or "extrato analitico" in t:
        return "fgts"
    if "trct" in t or "rescisao" in t or "termo de rescisao" in t:
        return "trct"
    if "laudo" in t or "pericial" in t or "ppp" in t:
        return "laudo_pericial"
    if "contrato" in t and "trabalho" in t:
        return "contrato"
    if "ctps" in t or "carteira de trabalho" in t:
        return "ctps"
    if "print" in t or "e-mail" in t or "email" in t or "audio" in t:
        return "print"
    return "outro"


def _achar_imagens_pra_tipo(
    tipo_canonico: str, imagens_pendentes: list
) -> list:
    """Remove e retorna todas as ImagemEmbedavel cujo tipo bate com o canonico.

    Mutates a lista pra remover as escolhidas (cada imagem casa com um item da
    lista de anexos, nao se repete). Match exato no campo .tipo (que veio
    sanitizado pelo Pydantic).
    """
    casadas = [img for img in imagens_pendentes if img.tipo == tipo_canonico]
    for img in casadas:
        imagens_pendentes.remove(img)
    return casadas


def _escrever_rol_documentos(
    doc: Any,
    anexos: Any,
    *,
    titulo: str = "VI — ROL DE DOCUMENTOS QUE INSTRUEM A PRESENTE",
    imagens_embedar: list | None = None,
) -> None:
    """Renderiza a secao ROL DE DOCUMENTOS com imagens embedded ou placeholders.

    Espera lista de dicts {numero, tipo, descricao}. Cada item vira:
    1. 'Doc. NN — TIPO: descricao' (justificado, com tipo em negrito)
    2. Imagem embedded (se houver match em imagens_embedar pelo tipo)
       OU '[ANEXAR ARQUIVO]' (centralizado, marcador visual fallback)

    PR15: `imagens_embedar` eh list[ImagemEmbedavel] do embed_processor.
    Quando presente, casa por tipo canonico (normalizado via _normalizar_tipo_anexo)
    e insere a imagem via doc.add_picture(). PDFs multi-pagina viram N imagens
    consecutivas. Imagens nao casadas com nenhum item ficam num bloco extra
    'OUTRAS PROVAS ANEXAS' no final.

    Itens malformados sao silenciosamente descartados.
    `titulo` permite numeracao romana dinamica no template builder.
    """
    if not isinstance(anexos, list) or not anexos:
        # Mesmo sem ROL textual, se houver imagens embedaveis ainda renderiza
        # o bloco de OUTRAS PROVAS ANEXAS no fim.
        if imagens_embedar:
            _escrever_outras_provas(doc, list(imagens_embedar))
        return

    itens_validos = [
        a for a in anexos
        if isinstance(a, dict) and (a.get("tipo") or a.get("descricao"))
    ]
    if not itens_validos:
        if imagens_embedar:
            _escrever_outras_provas(doc, list(imagens_embedar))
        return

    doc.add_heading(titulo, level=2)
    _add_paragraph(
        doc,
        "Acompanham a presente defesa os documentos abaixo relacionados, "
        "necessarios a comprovacao dos fatos impeditivos, modificativos e "
        "extintivos do direito do autor (art. 818, II, CLT):",
        justify=True,
    )

    # Copia mutavel — _achar_imagens_pra_tipo remove conforme casa.
    pendentes = list(imagens_embedar or [])

    for idx, item in enumerate(itens_validos[:10], start=1):
        numero = str(item.get("numero") or f"Doc. {idx:02d}").strip()
        tipo = str(item.get("tipo") or "").strip()
        descricao = str(item.get("descricao") or "").strip()

        para = doc.add_paragraph()
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_num = para.add_run(f"{numero} — ")
        run_num.bold = True
        if tipo:
            run_tipo = para.add_run(f"{tipo}")
            run_tipo.bold = True
        if descricao:
            sep = ": " if tipo else ""
            para.add_run(f"{sep}{descricao}")

        # PR15: tenta embedar imagens reais quando o tipo casa.
        tipo_canonico = _normalizar_tipo_anexo(tipo)
        imagens_casadas = _achar_imagens_pra_tipo(tipo_canonico, pendentes)

        if imagens_casadas:
            for img in imagens_casadas:
                _inserir_imagem(doc, img)
        else:
            # Sem imagem: mantem placeholder pro advogado anexar manualmente.
            placeholder = doc.add_paragraph()
            placeholder.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_ph = placeholder.add_run("[ANEXAR ARQUIVO]")
            run_ph.italic = True

    # Imagens que nao casaram com nenhum item — viram bloco extra "OUTRAS PROVAS".
    if pendentes:
        _escrever_outras_provas(doc, pendentes)


def _inserir_imagem(doc: Any, imagem: Any) -> None:
    """Insere uma ImagemEmbedavel centralizada com legenda opcional.

    Largura fixa em 15cm (cabe em folha A4 com margens padrao). Falha silenciosa:
    se python-docx nao conseguir abrir o blob, loga e segue.
    """
    try:
        para = doc.add_paragraph()
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(BytesIO(imagem.bytes_png), width=Cm(15))

        # Legenda discreta com nome do arquivo + numero da pagina (se multi).
        nome_legenda = imagem.nome
        if not imagem.eh_imagem_direta and imagem.pagina > 1:
            nome_legenda = f"{imagem.nome} (pag. {imagem.pagina})"
        elif not imagem.eh_imagem_direta:
            nome_legenda = f"{imagem.nome}"
        leg = doc.add_paragraph()
        leg.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        leg_run = leg.add_run(nome_legenda)
        leg_run.italic = True
        leg_run.font.size = Pt(9)
    except Exception as e:
        logger.warning("Falha ao inserir imagem %s no docx: %s", imagem.nome, e)
        # Fallback: placeholder text pra advogado anexar manualmente
        fb = doc.add_paragraph()
        fb.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fb_run = fb.add_run(f"[FALHA AO EMBEDAR {imagem.nome} — ANEXAR MANUALMENTE]")
        fb_run.italic = True


def _escrever_outras_provas(doc: Any, imagens: list) -> None:
    """Renderiza bloco final com imagens que nao casaram com itens do ROL."""
    if not imagens:
        return
    doc.add_paragraph("")
    doc.add_heading("OUTRAS PROVAS ANEXAS", level=3)
    for img in imagens:
        _inserir_imagem(doc, img)


def _strip_markdown(texto: Any) -> str:
    """Remove ## ### **bold** __underline__ pra texto plano legivel.

    Usado pelo programatico (sem template) — evita Markdown literal aparecendo
    como '## I - PRELIMINARMENTE' no .docx final.
    """
    if not texto:
        return ""
    s = str(texto)
    out = []
    for linha in s.split("\n"):
        stripped = linha.lstrip()
        if stripped.startswith("#"):
            out.append(stripped.lstrip("#").lstrip())
        else:
            out.append(linha)
    s = "\n".join(out)
    s = _RE_INLINE_BOLD.sub(r"\1", s)
    s = re.sub(r"__(.+?)__", r"\1", s)
    return s


def _escrever_secao_texto(doc: Any, titulo: str, conteudo: Any) -> None:
    """Adiciona heading + paragrafo justificado se conteudo for verdadeiro."""
    if not conteudo:
        return
    doc.add_heading(titulo, level=2)
    _add_paragraph(doc, _strip_markdown(conteudo), justify=True)


def _escrever_impugnacao_pedidos(doc: Any, impugnacoes: Any) -> None:
    """Renderiza impugnacao por pedido (dict) como SUBSECAO do merito.

    Antes era secao IV separada — agora aparece como III.Z dentro do
    merito, evitando duplicacao com as impugnacoes ja embutidas em III.A,
    III.B... do texto principal.
    """
    if not isinstance(impugnacoes, dict) or not impugnacoes:
        return
    doc.add_heading("III.Z — SINTESE DAS IMPUGNACOES ESPECIFICAS", level=3)
    for pedido, resposta in impugnacoes.items():
        paragrafo = doc.add_paragraph()
        run = paragrafo.add_run(f"Pedido: {pedido}")
        run.bold = True
        _add_paragraph(doc, _strip_markdown(resposta), justify=True)


def _escrever_rodape(doc: Any, minuta: dict[str, Any]) -> None:
    partes = [datetime.now().strftime("%Y-%m-%d %H:%M")]
    if minuta.get("observacoes"):
        partes.append(str(minuta["observacoes"]))
    _add_paragraph(doc, f"[{' | '.join(partes)}]")


def _add_paragraph(doc: Any, texto: str, *, justify: bool = False) -> None:
    paragrafo = doc.add_paragraph(texto)
    if justify:
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# ─────────────── Geracao .docx via modelo base do escritorio ──────────────


def montar_docx_com_modelo(
    modelo_b64: str,
    dados: dict[str, Any],
    minuta: dict[str, Any],
    *,
    imagens_embedar: list | None = None,
) -> bytes | None:
    """Preenche um .docx modelo do escritorio.

    Estrategia: abre o modelo como python-docx Document, limpa os paragrafos
    do body (mantem headers, footers, sections, watermark, sectPr), e insere
    paragrafos novos com a estrutura completa da contestacao no estilo do
    escritorio G. Trindade:
    - cabecalho com identificacao das partes (negrito em entidades-chave)
    - I — PRELIMINARMENTE com subsecoes A/B/C/D/E (negrito + sublinhado)
    - II — MERITO com subsecoes (cada uma ja contem a impugnacao do pedido)
    - III — DOS FUNDAMENTOS JURIDICOS
    - IV — DA AUTENTICIDADE DOS DOCUMENTOS
    - V — DOS PEDIDOS
    - numeracao 01.- 02.- ... global (renderer adiciona; IA so escreve texto puro)
    - Arial 11 justificado, aspas tipograficas curvas
    - travessao em dash (—) nos cabecalhos de secao
    - encerramento + assinatura

    PR15 — `imagens_embedar` (list[ImagemEmbedavel]) sao inseridas no ROL DE
    DOCUMENTOS final quando o tipo casa com algum item de documentos_anexos[].

    Retorna None em qualquer falha (caller faz fallback pro programatico).
    """
    modelo_bytes = _decodificar_modelo_b64(modelo_b64)
    if modelo_bytes is None:
        return None

    return _safe_step(
        "montar .docx a partir do modelo", _build_docx_from_template,
        modelo_bytes, dados, minuta, imagens_embedar,
    )


def _build_docx_from_template(
    modelo_bytes: bytes,
    dados: dict[str, Any],
    minuta: dict[str, Any],
    imagens_embedar: list | None = None,
) -> bytes:
    """Implementacao do montar_docx_com_modelo (separada pro _safe_step)."""
    doc = Document(BytesIO(modelo_bytes))

    # 0) ADAPTACAO DINAMICA: extrai fonte/espaco/tab do template ANTES de
    # limpar o body. Os helpers vao ler isso via ContextVar.
    estilo_extraido = _extrair_estilo_modelo(doc)
    token_estilo = _ESTILO_ATIVO.set(estilo_extraido)
    try:
        return _build_docx_from_template_inner(doc, dados, minuta, imagens_embedar)
    finally:
        _ESTILO_ATIVO.reset(token_estilo)


def _build_docx_from_template_inner(
    doc: Any,
    dados: dict[str, Any],
    minuta: dict[str, Any],
    imagens_embedar: list | None = None,
) -> bytes:
    """Lado pesado de _build_docx_from_template (apos extrair estilo)."""

    # 1) Limpa body preservando sectPr (config de pagina/secoes)
    body = doc.element.body
    sectPr = None
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            sectPr = deepcopy(child)
            continue
        body.remove(child)

    # 2) Contador global compartilhado entre todas as secoes numeradas
    contador = {"n": 0}

    # 3) CABECALHO PROCESSUAL (vem do Claude, com fallback para texto fixo se ausente)
    #    Claude entrega tudo em Markdown — renderer converte **bold** e __underline__.
    cabecalho_claude = minuta.get("cabecalho_processual")
    if cabecalho_claude:
        vara_str = dados.get("vara") or ""
        texto_cab = str(cabecalho_claude).strip()
        # Detecta separacao "Excelentissimo... <linha em branco> CONTEC..." e
        # renderiza como DOIS paragrafos distintos (alinhado ao modelo humano
        # do escritorio). Se Claude vier num bloco unico, faz split heuristico.
        partes = re.split(r"\n\s*\n", texto_cab, maxsplit=1)
        if len(partes) == 2 and "Excelent" in partes[0]:
            excelent_txt, partes_txt = partes[0].strip(), partes[1].strip()
        elif "Excelent" in texto_cab:
            # bloco unico — separa apos o primeiro ponto final
            m = re.search(r"(Excelent[^.]+\.)\s*(.*)", texto_cab, re.DOTALL)
            if m:
                excelent_txt, partes_txt = m.group(1).strip(), m.group(2).strip()
            else:
                excelent_txt, partes_txt = "", texto_cab
        else:
            excelent_txt = f"Excelentíssimo Senhor Doutor Juiz da {vara_str}."
            partes_txt = texto_cab
        if excelent_txt:
            _add_para_simples(doc, excelent_txt, negrito=True, alinhamento="justify")
            doc.add_paragraph("")
        if partes_txt:
            _add_text_para_com_indent(doc, partes_txt)
        doc.add_paragraph("")
    else:
        # Fallback: cabecalho default (caso o Claude nao gere — versao antiga)
        _add_para_simples(
            doc,
            f"Excelentíssimo Senhor Doutor Juiz da {dados.get('vara') or ''}.",
            negrito=True, alinhamento="justify",
        )
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _run(p, "\t")
        _run(p, f"{dados.get('reu') or ''}, ", bold=True)
        _run(p, "pessoa jurídica de direito privado, já qualificada nos autos da ")
        _run(p, f"Reclamação Trabalhista nº {dados.get('numero_processo') or ''}",
              bold=True, underline=True)
        _run(p, ", movida por ")
        _run(p, dados.get("autor") or "", bold=True)
        _run(p, ", vem respeitosamente à presença de V. Exa., por intermédio de "
                "seu advogado ao final assinado, apresentar ")
        _run(p, "CONTESTAÇÃO", bold=True)
        _run(p, " com fundamento no art. 847 da CLT, em conformidade com as "
                "razões expostas a seguir:")
        doc.add_paragraph("")

    # 4) Conteudo das secoes — numeracao romana DINAMICA conforme presenca de
    # secoes opcionais (Litigancia, Danos Morais, Fundamentos). I e II sao
    # fixos (Preliminares e Merito); demais sao renumeradas em sequencia.
    _ROMAN = ["", "I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX"]
    _add_secao(doc, minuta.get("preliminares"), contador, numerar=True,
               cabecalho_padrao=None)
    _add_secao(doc, minuta.get("merito"), contador, numerar=True,
               cabecalho_padrao="II — MERITO.")

    # impugnacao_pedidos NAO eh renderizado na peca final — fica como
    # referencia interna do prompt (keys snake_case nao saem para o cliente).
    # A IA ja deve incluir as impugnacoes nas subsecoes II.A, II.B... do merito.

    secao_idx = 3  # proxima apos II
    litigancia = minuta.get("litigancia_ma_fe")
    if litigancia and str(litigancia).strip():
        litigancia_norm = re.sub(
            r"^##\s*[IVXLCDM]+\s*[—\-]?\s*",
            f"## {_ROMAN[secao_idx]} — ",
            str(litigancia).strip(),
            count=1, flags=re.MULTILINE,
        )
        if not litigancia_norm.lstrip().startswith("##"):
            litigancia_norm = (
                f"## {_ROMAN[secao_idx]} — DA LITIGANCIA DE MA-FE.\n\n{litigancia_norm}"
            )
        _add_secao(doc, litigancia_norm, contador, numerar=True, cabecalho_padrao=None)
        secao_idx += 1

    danos = minuta.get("danos_morais")
    if danos and str(danos).strip():
        danos_norm = re.sub(
            r"^##\s*[IVXLCDM]+\s*[—\-]?\s*",
            f"## {_ROMAN[secao_idx]} — ",
            str(danos).strip(),
            count=1, flags=re.MULTILINE,
        )
        if not danos_norm.lstrip().startswith("##"):
            danos_norm = (
                f"## {_ROMAN[secao_idx]} — DA IMPROCEDENCIA DO PEDIDO DE INDENIZACAO POR DANOS MORAIS.\n\n{danos_norm}"
            )
        _add_secao(doc, danos_norm, contador, numerar=True, cabecalho_padrao=None)
        secao_idx += 1

    if minuta.get("fundamentos"):
        _add_secao(doc, minuta["fundamentos"], contador, numerar=True,
                   cabecalho_padrao=f"{_ROMAN[secao_idx]} — DOS FUNDAMENTOS JURIDICOS.")
        secao_idx += 1

    # AUTENTICIDADE (sempre presente)
    _add_heading(doc, f"{_ROMAN[secao_idx]} — DA AUTENTICIDADE DOS DOCUMENTOS.", nivel=1)
    doc.add_paragraph("")
    autenticidade_texto = minuta.get("autenticidade_documentos") or (
        "Nos termos do art. 830 da CLT, o patrono que subscreve a presente peça "
        "declara a autenticidade dos documentos acostados em cópia ao presente processo."
    )
    # Strip cabecalho que a IA as vezes prefixa por engano ('## V — DA AUTENTICIDADE...')
    autenticidade_texto = _strip_secao_prefix(str(autenticidade_texto))
    _add_text_para(doc, autenticidade_texto, contador, numerar=True)
    doc.add_paragraph("")
    secao_idx += 1

    # PEDIDOS
    _add_secao(doc, minuta.get("pedidos"), contador, numerar=False,
               cabecalho_padrao=f"{_ROMAN[secao_idx]} — DOS PEDIDOS.")
    secao_idx += 1

    # PR14 — ROL DE DOCUMENTOS PROBATORIOS A ANEXAR (logo apos PEDIDOS).
    # PR15 — Quando houver imagens_embedar, sao inseridas como imagens reais
    # no lugar do placeholder [ANEXAR ARQUIVO] (match por tipo canonico).
    anexos = minuta.get("documentos_anexos")
    if (isinstance(anexos, list) and anexos) or imagens_embedar:
        _escrever_rol_documentos(
            doc,
            anexos,
            titulo=f"{_ROMAN[secao_idx]} — ROL DE DOCUMENTOS QUE INSTRUEM A PRESENTE.",
            imagens_embedar=imagens_embedar,
        )
        doc.add_paragraph("")
        secao_idx += 1

    # Encerramento adicional (Claude tem que entregar com base legal correta)
    encerramento_texto = minuta.get("encerramento")
    if encerramento_texto:
        doc.add_paragraph("")
        _add_text_para(doc, str(encerramento_texto), contador, numerar=True)
        doc.add_paragraph("")

    protesta_texto = minuta.get("protesta_provas") or (
        "Protesta provar o alegado por todos os meios de prova em direito admitidos, "
        "em especial o depoimento pessoal do autor, sob pena de confissão, oitiva "
        "de testemunhas, perícia, juntada de documentos, dentre outros que se "
        "fizerem necessários para o esclarecimento dos fatos."
    )
    _add_text_para(doc, str(protesta_texto), contador, numerar=True)
    doc.add_paragraph("")

    # Encerramento formal: 'Termos em que, / p. deferimento.'
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Termos em que,")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "p. deferimento.")
    doc.add_paragraph("")

    # ASSINATURA (vem do Claude OU dados_advogado do form OU fallback)
    assinatura = minuta.get("assinatura") or {}
    if not isinstance(assinatura, dict):
        assinatura = {}
    def _limpar(s: Any) -> str:
        # Remove tabs/quebras espurias que vazam do JSON da IA (ex: 'OAB/PE a pree\tncher')
        return re.sub(r"\s+", " ", str(s or "")).strip()

    local_data = _limpar(assinatura.get("local_data"))
    # Se IA deixou 'data a preencher' ou similar, sobrescreve com fallback completo.
    if not local_data or "preencher" in local_data.lower() or "definir" in local_data.lower():
        local_data = f"{_local_padrao(dados)}, {_data_extenso()}."

    nome_adv = _limpar(assinatura.get("advogado"))
    if not nome_adv or "preencher" in nome_adv.lower() or "definir" in nome_adv.lower():
        nome_adv = _adv_padrao(dados, "nome")

    oab_adv = _limpar(assinatura.get("oab"))
    if not oab_adv or "preencher" in oab_adv.lower() or "definir" in oab_adv.lower():
        oab_adv = _adv_padrao(dados, "oab")

    _add_para_simples(doc, local_data, alinhamento="center")
    doc.add_paragraph("")
    doc.add_paragraph("")
    _add_para_simples(doc, nome_adv, alinhamento="center", negrito=True)
    _add_para_simples(doc, "Advogado", alinhamento="center")
    _add_para_simples(doc, oab_adv, alinhamento="center")

    # 5) Recoloca sectPr
    if sectPr is not None:
        body.append(sectPr)

    # 6) Salva
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


# ── Helpers de baixo nivel (criam paragraphs com a tipografia do escritorio) ─

def _font_default(run: Any) -> None:
    """Aplica fonte do template (default Arial 11).

    Le do contextvar _ESTILO_ATIVO que foi populado por _extrair_estilo_modelo
    quando o modelo base do escritorio foi inspecionado.
    """
    estilo = _estilo()
    run.font.name = estilo["font_name"]
    run.font.size = Pt(estilo["font_size_pt"])


def _aplicar_espacamento_padrao(p: Any) -> None:
    """Espacamento padrao via docx_style_defaults (line_spacing + space_after
    lidos do modelo com cap de seguranca aplicado em `_detectar_estilo_base`).

    A funcao centralizada em `docx_style_defaults.py` forca
    `WD_LINE_SPACING.MULTIPLE` explicitamente — sem isso o LibreOffice
    usa uma regra default que adiciona ar extra entre linhas, fazendo o
    mesmo DOCX renderizar com paginacao diferente entre Word e LO.
    """
    estilo = _estilo()
    _aplicar_espacamento_padrao_central(
        p,
        line_spacing=estilo["line_spacing"],
        space_after_pt=estilo["space_after_pt"],
        space_before_pt=0.0,
    )


def _aplicar_tab_stop_grande(p: Any) -> None:
    """Tab stop com o tamanho dominante do modelo (entre '01.-' e o texto)."""
    estilo = _estilo()
    pf = p.paragraph_format
    pf.tab_stops.add_tab_stop(Cm(estilo["tab_stop_cm"]))


def _run(
    p: Any,
    texto: str,
    *,
    bold: bool = False,
    underline: bool = False,
    italic: bool = False,
) -> None:
    """Adiciona um run a `p` com formatacao Arial 11 + negrito/sublinhado/italico."""
    r = p.add_run(texto)
    r.bold = bold
    r.underline = underline
    r.italic = italic
    _font_default(r)


def _add_para_simples(
    doc: Any,
    texto: str,
    *,
    negrito: bool = False,
    alinhamento: str = "justify",
) -> None:
    """Adiciona um paragrafo simples (sem Markdown) com alinhamento."""
    p = doc.add_paragraph()
    if alinhamento == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alinhamento == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _run(p, texto, bold=negrito)


def _add_heading(doc: Any, texto: str, *, nivel: int) -> None:
    """Adiciona cabecalho com negrito + sublinhado.

    Espacamento antes e line_spacing lidos do template:
    - nivel 1 (I, II, III...): respiracao maior (default 12pt)
    - nivel 2/3 (A, B, C...):  respiracao menor (default 8pt)

    Usa o helper centralizado de `docx_style_defaults` (mesma regra
    WD_LINE_SPACING.MULTIPLE explicita pra garantir paginacao identica
    em Word e LibreOffice).
    """
    estilo = _estilo()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _aplicar_espacamento_padrao_central(
        p,
        line_spacing=estilo["line_spacing"],
        space_after_pt=estilo["space_after_pt"],
        space_before_pt=(
            estilo["space_before_secao1_pt"]
            if nivel == 1
            else estilo["space_before_secao2_pt"]
        ),
    )
    texto_limpo = _RE_INLINE_BOLD.sub(r"\1", texto)
    if nivel == 1:
        texto_limpo = texto_limpo.upper()
    _run(p, texto_limpo, bold=True, underline=True)


def _add_text_para(
    doc: Any, texto: str, contador: dict[str, int], *, numerar: bool
) -> None:
    """Adiciona paragrafo de texto justificado com numeracao opcional, bold e underline inline.

    Suporta:
    - '**texto**' -> run em negrito
    - '__texto__' -> run em negrito + sublinhado (estilo G. Trindade pra processo)
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _aplicar_espacamento_padrao(p)
    texto_join = " ".join(line.strip() for line in texto.split("\n") if line.strip())
    # Strip prefixos 'NN.-' / 'NN.- NN.- ' que a IA as vezes inclui no inicio do
    # texto. Evita duplicacao com a numeracao automatica do renderer.
    texto_join = re.sub(r"^(?:\d{1,3}[\.\-]+\s+)+", "", texto_join).strip()
    if numerar:
        _aplicar_tab_stop_grande(p)
        contador["n"] += 1
        # Padrao do modelo G. Trindade: 'NN.- TAB ESPACO TAB' produz o espaco
        # largo entre o numero do paragrafo e o texto. Reproduzimos identico.
        _run(p, f"{contador['n']:02d}.-\t \t")
    _render_inline_markdown(p, texto_join)


def _add_secao(
    doc: Any,
    texto: Any,
    contador: dict[str, int],
    *,
    numerar: bool,
    cabecalho_padrao: str | None,
) -> None:
    """Renderiza uma secao do Claude (texto Markdown) como sequencia de paragrafos.

    Se o texto NAO comeca com '##' ou '#', adiciona o cabecalho_padrao antes.
    Linhas '##' viram heading 1, '###' viram heading 2.
    Blocos onde TODAS linhas comecam com '> ' viram quote (italico+recuado).
    Outros viram texto numerado.
    """
    if not texto:
        if cabecalho_padrao:
            _add_heading(doc, cabecalho_padrao, nivel=1)
            doc.add_paragraph("")
        return

    s = str(texto).strip()

    primeira = s.split("\n", 1)[0].strip()
    if cabecalho_padrao and not primeira.startswith("#"):
        _add_heading(doc, cabecalho_padrao, nivel=1)
        doc.add_paragraph("")

    blocos = re.split(r"\n\s*\n", s)
    for bloco in blocos:
        bloco = bloco.strip()
        if not bloco:
            continue
        primeira_linha = bloco.split("\n", 1)[0].strip()
        m1 = _RE_HEADING1.match(primeira_linha)
        m2 = _RE_HEADING2.match(primeira_linha)

        # Detecta blockquote: todas as linhas (nao vazias) comecam com '>'
        linhas_bloco = [l for l in bloco.split("\n") if l.strip()]
        is_quote = all(_RE_QUOTE_LINE.match(l) for l in linhas_bloco) if linhas_bloco else False

        if is_quote:
            _add_quote_para(doc, bloco)
            doc.add_paragraph("")
        elif m2:
            _add_heading(doc, m2.group(1), nivel=2)
            doc.add_paragraph("")
            resto = bloco.split("\n", 1)[1].strip() if "\n" in bloco else ""
            if resto:
                _add_text_para(doc, resto, contador, numerar=numerar)
                doc.add_paragraph("")
        elif m1:
            _add_heading(doc, m1.group(1), nivel=1)
            doc.add_paragraph("")
            resto = bloco.split("\n", 1)[1].strip() if "\n" in bloco else ""
            if resto:
                _add_text_para(doc, resto, contador, numerar=numerar)
                doc.add_paragraph("")
        else:
            _add_text_para(doc, bloco, contador, numerar=numerar)
            doc.add_paragraph("")


def _add_quote_para(doc: Any, bloco: str) -> None:
    """Renderiza um bloco '> texto' como paragrafo recuado em italico.

    Mimetiza o estilo do modelo G. Trindade pra citacoes de lei/sumula.
    """
    linhas = []
    for l in bloco.split("\n"):
        m = _RE_QUOTE_LINE.match(l)
        if m:
            linhas.append(m.group(1).strip())
    texto = " ".join(l for l in linhas if l)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Recuo do blockquote lido do modelo (default 3.0cm)
    p.paragraph_format.left_indent = Cm(_estilo()["quote_left_indent_cm"])
    p.paragraph_format.right_indent = Cm(0.5)
    _aplicar_espacamento_padrao(p)
    # Texto em italico — bold inline ainda funciona dentro do quote
    _RE_COMBO = re.compile(r"\*\*(.+?)\*\*|__(.+?)__")
    pos = 0
    for m in _RE_COMBO.finditer(texto):
        if m.start() > pos:
            _run(p, texto[pos:m.start()], italic=True)
        if m.group(1) is not None:
            _run(p, m.group(1), bold=True, italic=True)
        else:
            _run(p, m.group(2), bold=True, underline=True, italic=True)
        pos = m.end()
    if pos < len(texto):
        _run(p, texto[pos:], italic=True)


def _add_secao_dict(
    doc: Any,
    impugnacoes: Any,
    contador: dict[str, int],
    *,
    cabecalho: str,
) -> None:
    """Renderiza o dict de impugnacao_pedidos com cabecalho + items numerados."""
    _add_heading(doc, cabecalho, nivel=1)
    doc.add_paragraph("")
    if isinstance(impugnacoes, dict):
        for pedido, resposta in impugnacoes.items():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            contador["n"] += 1
            _run(p, f"{contador['n']:02d}.-\t")
            _run(p, f"{pedido}: ", bold=True)
            # Limpa Markdown da resposta e aplica bold inline
            resp_txt = " ".join(
                line.strip() for line in str(resposta).split("\n") if line.strip()
            )
            pos = 0
            for m in _RE_INLINE_BOLD.finditer(resp_txt):
                if m.start() > pos:
                    _run(p, resp_txt[pos:m.start()])
                _run(p, m.group(1), bold=True)
                pos = m.end()
            if pos < len(resp_txt):
                _run(p, resp_txt[pos:])
            doc.add_paragraph("")
    elif impugnacoes:
        _add_text_para(doc, str(impugnacoes), contador, numerar=True)
        doc.add_paragraph("")


def _decodificar_modelo_b64(modelo_b64: str) -> bytes | None:
    try:
        return base64.b64decode(modelo_b64.strip(), validate=True)
    except (binascii.Error, ValueError) as error:
        logger.warning("Modelo base nao decodificou como base64: %s", error)
        return None


def _safe_step(descricao: str, func: Callable, *args, **kwargs):
    """Roda `func(*args, **kwargs)` mapeando excecoes do docx para None+log.

    Centraliza o tratamento uniforme das falhas possiveis no `_build_docx_from_template`
    (abrir bytes, manipular XML do body, salvar) sem repetir try/except.

    `python-docx` levanta varios subtipos (`PackageNotFoundError`, `KeyError` em
    runs partidos, `OSError`). Capturamos `Exception` mas registramos o nome
    real da excecao em log — diagnostico mostra o tipo real ao inves de
    `Exception` generico.
    """
    try:
        return func(*args, **kwargs)
    except Exception as error:  # noqa: BLE001 - lib externa, fallback obrigatorio
        logger.warning(
            "Falha ao %s: %s: %s", descricao, type(error).__name__, error
        )
        return None


def _safe_void_step(descricao: str, func: Callable, *args, **kwargs) -> bool:
    """Como _safe_step, mas para chamadas que NAO retornam valor.

    Funcoes como `python-docx` `doc.save()` retornam None em sucesso. Usar
    `_safe_step` com elas resulta em falso negativo (None confunde com falha).
    Este helper retorna True em sucesso e False em excecao.
    """
    try:
        func(*args, **kwargs)
        return True
    except Exception as error:  # noqa: BLE001 - lib externa
        logger.warning(
            "Falha ao %s: %s: %s", descricao, type(error).__name__, error
        )
        return False


# ── Regex compartilhadas pelo renderer python-docx ───────────────────────────
_RE_HEADING1 = re.compile(r"^#{1,2}\s+(.+?)\s*$")  # ## TITULO  ou  # TITULO
_RE_HEADING2 = re.compile(r"^#{3,4}\s+(.+?)\s*$")  # ### Subtitulo
_RE_INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
_RE_INLINE_UNDER = re.compile(r"__(.+?)__")
_RE_QUOTE_LINE = re.compile(r"^\s*>\s?(.*)$")

_MESES_PT = {
    "January": "janeiro", "February": "fevereiro", "March": "março",
    "April": "abril", "May": "maio", "June": "junho", "July": "julho",
    "August": "agosto", "September": "setembro", "October": "outubro",
    "November": "novembro", "December": "dezembro",
}


def _data_extenso() -> str:
    """Retorna data atual no formato 'DD de mes de AAAA' (portugues)."""
    s = datetime.now().strftime("%d de %B de %Y")
    for en, pt in _MESES_PT.items():
        s = s.replace(en, pt)
    return s


def _local_padrao(dados: dict[str, Any]) -> str:
    """Local padrao para encerramento: sede do escritorio G. Trindade (Recife/PE).

    A cidade da Vara nao deve aparecer na assinatura (so no cabecalho processual);
    no encerramento usa-se a sede do escritorio.
    """
    return "Recife/PE"


def _adv_padrao(dados: dict[str, Any], campo: str) -> str:
    """Pega dados do advogado do form OU fallback para o escritorio G. Trindade."""
    adv = dados.get("dados_advogado") or {}
    if not isinstance(adv, dict):
        adv = {}
    if campo == "nome":
        return adv.get("nome") or "Genner Trindade"
    if campo == "oab":
        return adv.get("oab") or "OAB/PE 27.790"
    return ""


def _strip_secao_prefix(texto: str) -> str:
    """Remove prefixos de cabecalho de secao gerados pela IA dentro do conteudo.

    A IA as vezes prefixa '## V — DA AUTENTICIDADE...' + 'NN.- ' no inicio do
    campo (que ja eh renderizado com cabecalho proprio pelo builder). Sem strip,
    o cabecalho duplica e a numeracao vaza como texto.
    """
    if not texto:
        return texto
    s = str(texto).lstrip()
    # Remove '## I-VI — CABECALHO...' no inicio (com ou sem ponto final)
    s = re.sub(r"^#{1,4}\s*[IVXLCDM]+\s*[—\-–]\s*[^\n]+?\n+", "", s, count=1)
    # Remove '## CABECALHO SEM ROMANO' no inicio tambem
    s = re.sub(r"^#{1,4}\s*[^\n]+?\n+", "", s, count=1)
    # Remove numeracao 'NN.-' duplicada que a IA as vezes adiciona
    s = re.sub(r"^\d{1,3}[\.\-]+\s+", "", s.lstrip())
    return s.strip()


def _add_text_para_com_indent(doc: Any, texto: str) -> None:
    """Paragrafo de identificacao processual: indentado com TAB e bold inline.

    Texto vem do Claude no formato 'TEXTO **negrito** e __sublinhado__'.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _aplicar_espacamento_padrao(p)
    _run(p, "\t")
    texto_join = " ".join(line.strip() for line in texto.split("\n") if line.strip())
    _render_inline_markdown(p, texto_join)


def _tipograficar_aspas(texto: str) -> str:
    """Substitui aspas duplas planas por curvas tipograficas (estilo G. Trindade).

    So toca aspas DUPLAS (\" → “/”). Apostrofes (') sao deixadas intactas —
    em PT-BR juridico sao raras e o pareamento heuristico falharia em
    contracoes isoladas (ex: 'd'agua', 'p'ra').

    Heuristica baseada em PAREAMENTO: a 1a \" abre (“), a 2a fecha (”),
    alternando. Reseta o estado a cada quebra de linha para que aspa nao
    fechada num paragrafo nao contamine o seguinte. Idempotente — aspas
    tipograficas pre-existentes nao sao tocadas.
    """
    if not texto or '"' not in texto:
        return texto
    out = []
    in_double = False
    for ch in texto:
        if ch == "\n":
            in_double = False
            out.append(ch)
        elif ch == '"':
            out.append("”" if in_double else "“")
            in_double = not in_double
        else:
            out.append(ch)
    return "".join(out)


def _render_inline_markdown(p: Any, texto: str) -> None:
    """Renderiza **bold** e __underline__ inline num paragrafo existente.

    Processa o texto sequencialmente: cada match de **...** vira run em
    negrito, __...__ vira run sublinhado, resto fica como texto plano.
    Aspas planas sao convertidas em tipograficas curvas (estilo escritorio).
    """
    texto = _tipograficar_aspas(texto)
    # Combina os 2 regex em um soh: alterna marcadores
    _RE_COMBO = re.compile(r"\*\*(.+?)\*\*|__(.+?)__")
    pos = 0
    for m in _RE_COMBO.finditer(texto):
        if m.start() > pos:
            _run(p, texto[pos:m.start()])
        if m.group(1) is not None:  # **bold**
            _run(p, m.group(1), bold=True)
        else:  # __underline__
            _run(p, m.group(2), bold=True, underline=True)
        pos = m.end()
    if pos < len(texto):
        _run(p, texto[pos:])
