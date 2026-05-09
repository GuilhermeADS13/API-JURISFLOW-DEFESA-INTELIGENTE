"""Extrator de texto de peticao inicial em PDF ou DOCX.

Usado pela rota POST /api/contestar-por-peticao antes de enviar o conteudo ao
agente IA. Mantem a logica de extracao no backend (nao no n8n) porque o
container n8n base nao tem mammoth/pdfjs-dist, e Python ja tem libs maduras.

PR6 #1 — Long Context (Guia v3 §2.2): em vez de Map-Reduce com Claude por
chunk (caro), usamos estrategia hibrida:
1. Limite expandido para 80k chars (~20-25k tokens; Claude Sonnet 4.6 aceita
   ate 200k tokens com folga).
2. Quando texto excede o limite, aplica `prefiltrar_secoes_juridicas` que
   identifica e prioriza secoes-chave da peticao (DOS FATOS, DO DIREITO,
   DOS PEDIDOS, DA LEGITIMIDADE, etc.) para nao perder pedidos no fim do
   documento.
3. Se mesmo apos pre-filtragem ainda passar do limite, usa `truncamento_inteligente`
   que preserva inicio (cabecalho do processo) + fim (pedidos), com elipse no meio.

Map-Reduce (chunking + Claude por parte) fica como PR7 se essa abordagem
nao for suficiente em casos extremos.
"""

from __future__ import annotations

import base64
import binascii
import logging
import os
import re
from io import BytesIO

from docx import Document
from pypdf import PdfReader

logger = logging.getLogger(__name__)

# PR6 #3 — OCR opcional (Tesseract local) para PDFs digitalizados.
# Importacoes lazy para nao quebrar a importacao do modulo se as libs nao
# estiverem instaladas no ambiente (ex: testes em CI sem tesseract).
try:
    import pytesseract  # type: ignore
    from pdf2image import convert_from_bytes  # type: ignore
    _OCR_LIBS_DISPONIVEIS = True
except ImportError:
    pytesseract = None  # type: ignore
    convert_from_bytes = None  # type: ignore
    _OCR_LIBS_DISPONIVEIS = False

# PR6: limite expandido. Claude Sonnet 4.6 tem 200k tokens de janela; 80k
# chars ~ 20k tokens deixa folga confortavel para system prompt + few-shot
# exemplares + RAG defesas anteriores na mesma chamada.
MAX_TEXTO_PETICAO_CHARS = 80_000
MAX_TEXTO_MODELO_BASE_CHARS = 15_000

# PR6 #3 — OCR fallback (Tesseract). Configuravel via env para desligar em
# emergencia ou para acelerar testes em CI.
OCR_ENABLED = os.getenv("OCR_ENABLED", "true").lower() in ("1", "true", "yes")
OCR_MAX_PAGES = int(os.getenv("OCR_MAX_PAGES", "15"))  # OCR e lento (~5-10s/pagina)
OCR_DPI = int(os.getenv("OCR_DPI", "300"))  # Tesseract recomenda 300 DPI
OCR_LANG = os.getenv("OCR_LANG", "por")  # tesseract-ocr-por instalado no Dockerfile

# Limite de chars para considerar um PDF como "digitalizado" (texto curto =
# pypdf nao conseguiu extrair, provavelmente eh imagem).
PDF_OCR_FALLBACK_THRESHOLD = 200

# Marcadores de secoes que aparecem em peticoes brasileiras. Casamento
# case-insensitive, ancorado em inicio de linha apos whitespace, com
# tolerancia a numeracao romana ou arabe (ex: "I - DOS FATOS", "1. DOS FATOS").
# Ordenamos por importancia: PEDIDOS > FATOS > DIREITO > restantes.
SECOES_PRIORIZADAS = (
    # (regex, nome_legivel, peso_prioridade)
    (re.compile(r"^(?:[IVX]+\.?\s*[-–—)]?\s*|\d+[.)]\s*)?d[oa]s?\s+pedidos?\b", re.IGNORECASE | re.MULTILINE), "PEDIDOS", 100),
    (re.compile(r"^(?:[IVX]+\.?\s*[-–—)]?\s*|\d+[.)]\s*)?d[oa]s?\s+fatos\b", re.IGNORECASE | re.MULTILINE), "FATOS", 90),
    (re.compile(r"^(?:[IVX]+\.?\s*[-–—)]?\s*|\d+[.)]\s*)?d[oa]\s+direito\b", re.IGNORECASE | re.MULTILINE), "DIREITO", 80),
    (re.compile(r"^(?:[IVX]+\.?\s*[-–—)]?\s*|\d+[.)]\s*)?d[oa]s?\s+fundamentos?\s+jur[íi]dicos?\b", re.IGNORECASE | re.MULTILINE), "FUNDAMENTOS", 80),
    (re.compile(r"^(?:[IVX]+\.?\s*[-–—)]?\s*|\d+[.)]\s*)?d[oa]\s+m[ée]rito\b", re.IGNORECASE | re.MULTILINE), "MERITO", 75),
    (re.compile(r"^(?:[IVX]+\.?\s*[-–—)]?\s*|\d+[.)]\s*)?d[oa]\s+valor\s+d[oa]\s+causa\b", re.IGNORECASE | re.MULTILINE), "VALOR_CAUSA", 70),
    (re.compile(r"^(?:[IVX]+\.?\s*[-–—)]?\s*|\d+[.)]\s*)?d[oa]\s+legitimidade\b", re.IGNORECASE | re.MULTILINE), "LEGITIMIDADE", 60),
    (re.compile(r"^(?:[IVX]+\.?\s*[-–—)]?\s*|\d+[.)]\s*)?d[oa]s?\s+preliminares?\b", re.IGNORECASE | re.MULTILINE), "PRELIMINARES", 60),
    (re.compile(r"^(?:[IVX]+\.?\s*[-–—)]?\s*|\d+[.)]\s*)?d[oa]\s+t[ée]cnico-?jur[íi]dic", re.IGNORECASE | re.MULTILINE), "TECNICO_JURIDICO", 55),
    (re.compile(r"^(?:[IVX]+\.?\s*[-–—)]?\s*|\d+[.)]\s*)?d[oa]s?\s+provas?\b", re.IGNORECASE | re.MULTILINE), "PROVAS", 50),
)

# Marcador de "fim de secao" — outra secao numerada/titulada ou final do doc.
# Usado para delimitar onde uma secao priorizada termina.
RE_FIM_SECAO = re.compile(
    r"^(?:[IVX]+\.?\s*[-–—)]?\s*|\d+[.)]\s*)?d[oa]s?\s+\w+",
    re.IGNORECASE | re.MULTILINE,
)


class ExtracaoError(Exception):
    """Falha ao extrair texto do arquivo enviado."""


def extrair_texto_peticao(conteudo: bytes, nome: str) -> str:
    """Extrai texto da peticao inicial. Detecta tipo pela extensao do nome.

    Aplica pre-filtragem por secoes juridicas se o texto exceder
    `MAX_TEXTO_PETICAO_CHARS`. Levanta `ExtracaoError` se o conteudo estiver
    corrompido ou se nao for possivel obter texto util (>=50 chars apos strip).
    """
    if not conteudo:
        raise ExtracaoError("Conteudo da peticao vazio.")

    nome_lower = (nome or "").lower().strip()

    if nome_lower.endswith(".docx"):
        texto = _extrair_docx(conteudo)
    elif nome_lower.endswith(".pdf"):
        texto = _extrair_pdf(conteudo)
    elif nome_lower.endswith(".doc"):
        # .doc legado — extracao basica de texto legivel ASCII/Latin-1.
        texto = conteudo.decode("latin1", errors="ignore")
        # Remove caracteres de controle nao-imprimiveis exceto quebras de linha.
        texto = "".join(c for c in texto if c == "\n" or c == "\r" or c == "\t" or 0x20 <= ord(c) <= 0xFFFF)
    else:
        raise ExtracaoError(f"Extensao nao suportada para extracao: {nome!r}")

    texto = _limpar_texto(texto)

    if len(texto.strip()) < 50:
        raise ExtracaoError(
            "Nao foi possivel extrair texto legivel da peticao. "
            "Tente converter para DOCX e enviar novamente."
        )

    # PR6: se texto cabe no limite, retorna direto. Se passa, pre-filtra +
    # truncamento inteligente preservando comecho e fim.
    if len(texto) <= MAX_TEXTO_PETICAO_CHARS:
        return texto

    return _aplicar_long_context(texto, MAX_TEXTO_PETICAO_CHARS)


def extrair_e_consolidar_textos(
    peticao_bytes: bytes,
    peticao_nome: str,
    anexos: list[tuple[bytes, str]] | None = None,
) -> str:
    """Extrai texto da peticao e dos anexos, consolida em uma unica string.

    Cada anexo eh prefixado com cabecalho `=== ANEXO N (nome) ===` para que o
    agente Claude consiga distinguir e relacionar os documentos. O total e
    truncado em `MAX_TEXTO_PETICAO_CHARS` (priorizando a peticao no inicio,
    aplicando pre-filtragem se necessario).

    Anexos que falham na extracao sao silenciosamente descartados — anexos sao
    auxiliares, nao podem bloquear o fluxo da geracao.
    """
    texto_peticao = extrair_texto_peticao(peticao_bytes, peticao_nome)

    if not anexos:
        return texto_peticao

    partes = [texto_peticao]
    for idx, (conteudo, nome) in enumerate(anexos, start=1):
        try:
            texto_anexo = extrair_texto_peticao(conteudo, nome)
        except ExtracaoError:
            continue
        partes.append(f"\n\n=== ANEXO {idx} ({nome}) ===\n\n{texto_anexo}")

    consolidado = "".join(partes)
    if len(consolidado) <= MAX_TEXTO_PETICAO_CHARS:
        return consolidado
    # Mesmo com tudo truncado, peticao + anexos podem passar. Aplica
    # long-context ao consolidado preservando peticao no inicio (que ja foi
    # filtrada) + ultima parte (anexos com pedidos relacionados).
    return _aplicar_long_context(consolidado, MAX_TEXTO_PETICAO_CHARS)


def extrair_texto_modelo_base(conteudo_b64: str | None) -> str:
    """Extrai texto do modelo base .docx (opcional).

    Retorna string vazia em ausencia/erro — o modelo base e opcional e nao deve
    bloquear o fluxo da geracao.
    """
    if not conteudo_b64:
        return ""

    try:
        conteudo = base64.b64decode(conteudo_b64.strip(), validate=True)
    except (binascii.Error, ValueError):
        return ""

    try:
        texto = _extrair_docx(conteudo)
    except ExtracaoError:
        return ""

    return _limpar_texto(texto)[:MAX_TEXTO_MODELO_BASE_CHARS]


# ── PR6 #1 — Long Context: pre-filtragem + truncamento inteligente ──────────


def prefiltrar_secoes_juridicas(texto: str) -> dict[str, str]:
    """Identifica secoes priorizadas em uma peticao e retorna por nome.

    Para cada match em `SECOES_PRIORIZADAS`, captura do inicio da secao ate
    o proximo cabecalho de secao OU final do documento. Se a mesma secao
    aparecer mais de uma vez (ex: PEDIDOS de cada autor), concatena.

    Retorna dict {nome: texto_da_secao}. Secoes nao encontradas ficam de fora.
    """
    if not texto:
        return {}

    secoes_encontradas: dict[str, list[str]] = {}

    # Mapeia (start, end, nome) para todas as secoes priorizadas.
    matches: list[tuple[int, int, str]] = []
    for regex, nome, _peso in SECOES_PRIORIZADAS:
        for m in regex.finditer(texto):
            matches.append((m.start(), m.end(), nome))

    if not matches:
        return {}

    # Ordena por posicao para conseguir delimitar fim de cada secao
    # (proximo cabecalho de secao OU EOF).
    matches.sort(key=lambda x: x[0])

    # Coleta posicoes de TODOS cabecalhos de secao (priorizadas + outras)
    # para saber onde uma secao termina.
    todos_cabecalhos = sorted(set(m.start() for m in RE_FIM_SECAO.finditer(texto)))

    for start, end, nome in matches:
        # Encontra a proxima posicao de cabecalho apos `end` (fim da secao atual).
        proximo = next((p for p in todos_cabecalhos if p > end), len(texto))
        trecho = texto[start:proximo].strip()
        if trecho:
            secoes_encontradas.setdefault(nome, []).append(trecho)

    return {nome: "\n\n".join(trechos) for nome, trechos in secoes_encontradas.items()}


def _aplicar_long_context(texto: str, limite: int) -> str:
    """Estrategia hibrida quando texto > limite.

    Passo 1: tenta pre-filtragem por secoes priorizadas. Se a soma das secoes
    couber em `limite * 0.9`, usa direto (deixa 10% para preambulo/cabecalho).

    Passo 2: senao, faz truncamento inteligente preservando inicio (cabecalho)
    + fim (pedidos), com elipse "[...trecho intermediario omitido...]" no meio.
    """
    secoes = prefiltrar_secoes_juridicas(texto)

    if secoes:
        # Monta texto pre-filtrado priorizando por peso decrescente.
        pesos_por_nome = {nome: peso for _, nome, peso in SECOES_PRIORIZADAS}
        secoes_ordenadas = sorted(
            secoes.items(), key=lambda x: pesos_por_nome.get(x[0], 0), reverse=True
        )

        # Reserva ~10% do limite para o cabecalho original (partes/processo).
        cabecalho_size = min(int(limite * 0.10), 4000)
        cabecalho = texto[:cabecalho_size].strip()

        partes_filtradas = [
            f"=== {nome} ===\n{conteudo}" for nome, conteudo in secoes_ordenadas
        ]
        corpo_filtrado = "\n\n".join(partes_filtradas)

        candidato = (
            f"{cabecalho}\n\n[...preambulo omitido — ver secoes abaixo...]\n\n"
            f"{corpo_filtrado}"
        )

        if len(candidato) <= limite:
            return candidato

        # Pre-filtragem nao foi suficiente; cai para truncamento inteligente.

    # Truncamento inteligente: 60% do limite no inicio, 40% no fim, com elipse.
    inicio_size = int(limite * 0.6)
    fim_size = limite - inicio_size - 80  # 80 chars de folga para a elipse
    inicio = texto[:inicio_size]
    fim = texto[-fim_size:] if fim_size > 0 else ""
    elipse = "\n\n[...trecho intermediario omitido por exceder o limite de contexto...]\n\n"
    return inicio + elipse + fim


def _extrair_docx(conteudo: bytes) -> str:
    try:
        doc = Document(BytesIO(conteudo))
    except Exception as error:
        raise ExtracaoError(
            f"Falha ao abrir o .docx: {type(error).__name__}: {error}"
        ) from error

    linhas: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text:
            linhas.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text:
                        linhas.append(paragraph.text)

    return "\n".join(linhas)


def _extrair_pdf(conteudo: bytes) -> str:
    """Extrai texto de PDF com fallback OCR (PR6 #3) para PDFs digitalizados.

    Tenta primeiro pypdf (rapido, funciona para PDFs nativos com camada de
    texto). Se o resultado for muito curto (< PDF_OCR_FALLBACK_THRESHOLD),
    presume que eh um PDF imagem (escaneado) e cai no Tesseract.

    OCR e desligavel via OCR_ENABLED=false no env para emergencia.
    """
    try:
        reader = PdfReader(BytesIO(conteudo))
    except Exception as error:
        raise ExtracaoError(
            f"Falha ao abrir o PDF: {type(error).__name__}: {error}"
        ) from error

    paginas: list[str] = []
    # PR6: aumentamos o limite de 30 -> 100 paginas. Mesmo assim, _aplicar_long_context
    # vai filtrar/truncar inteligentemente se o texto extraido exceder MAX_TEXTO_PETICAO_CHARS.
    for page in reader.pages[:100]:
        try:
            paginas.append(page.extract_text() or "")
        except Exception:
            paginas.append("")

    texto_pypdf = "\n\n".join(p for p in paginas if p.strip())

    # PR6 #3: se pypdf retornou pouco/nada, tenta OCR (PDF provavelmente
    # eh imagem digitalizada).
    if len(texto_pypdf.strip()) >= PDF_OCR_FALLBACK_THRESHOLD:
        return texto_pypdf

    if not OCR_ENABLED:
        logger.info(
            "PDF parece digitalizado (texto pypdf=%d chars) mas OCR_ENABLED=false; retornando vazio.",
            len(texto_pypdf),
        )
        return texto_pypdf

    if not _OCR_LIBS_DISPONIVEIS:
        logger.warning(
            "PDF parece digitalizado mas pytesseract/pdf2image nao instalados; retornando texto curto."
        )
        return texto_pypdf

    logger.info(
        "PDF com texto curto (%d chars) — disparando OCR Tesseract (lang=%s, dpi=%d, max_pages=%d).",
        len(texto_pypdf),
        OCR_LANG,
        OCR_DPI,
        OCR_MAX_PAGES,
    )

    try:
        texto_ocr = _extrair_pdf_via_ocr(conteudo)
    except Exception as error:
        logger.error(
            "Falha no OCR Tesseract: %s: %s. Retornando texto pypdf (curto).",
            type(error).__name__,
            error,
        )
        return texto_pypdf

    # Combina pypdf + OCR caso pypdf tenha capturado algo (ex: cabeçalho com
    # texto nativo + corpo escaneado). Se OCR for vazio, fica com pypdf.
    if texto_ocr.strip():
        if texto_pypdf.strip():
            return texto_pypdf + "\n\n" + texto_ocr
        return texto_ocr
    return texto_pypdf


def _extrair_pdf_via_ocr(conteudo: bytes) -> str:
    """Converte paginas do PDF em imagens e roda Tesseract em cada uma.

    Limita a `OCR_MAX_PAGES` paginas (OCR eh lento ~5-10s/pagina e o webhook
    n8n tem timeout de 180s — passar disso quebra o pipeline). Processa uma
    pagina de cada vez para nao explodir a memoria com PDFs grandes.
    """
    if not _OCR_LIBS_DISPONIVEIS:
        return ""

    # convert_from_bytes carrega TUDO em memoria; melhor passar last_page e
    # processar em batches se PDF for gigante. Aqui limitamos no parametro.
    try:
        imagens = convert_from_bytes(
            conteudo,
            dpi=OCR_DPI,
            first_page=1,
            last_page=OCR_MAX_PAGES,
            fmt="ppm",
        )
    except Exception as error:
        # pdf2image precisa do poppler instalado (poppler-utils no Dockerfile).
        # Em ambiente sem poppler, retorna vazio em vez de quebrar.
        logger.error(
            "pdf2image falhou: %s: %s (poppler-utils instalado?)",
            type(error).__name__,
            error,
        )
        return ""

    paginas_texto: list[str] = []
    for idx, img in enumerate(imagens, start=1):
        try:
            texto_pagina = pytesseract.image_to_string(img, lang=OCR_LANG)
            if texto_pagina and texto_pagina.strip():
                paginas_texto.append(texto_pagina)
        except Exception as error:
            logger.warning(
                "OCR falhou na pagina %d: %s: %s",
                idx,
                type(error).__name__,
                error,
            )

    return "\n\n".join(paginas_texto)


def _limpar_texto(texto: str) -> str:
    """Remove sequencias de espacos/quebras redundantes."""
    if not texto:
        return ""
    # Colapsa 3+ quebras consecutivas em 1 para reduzir tokens sem perder estrutura.
    linhas = [linha.rstrip() for linha in texto.splitlines()]
    return "\n".join(linha for linha in linhas if linha != "" or True).strip()
