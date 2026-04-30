"""Editor cirurgico de .docx que preserva formatacao quando possivel.

Estrategia:
- Se a string `antigo` cabe inteira em um unico run, substitui no proprio run
  e preserva todo o estilo (negrito, italico, fonte, cor).
- Se a string `antigo` cruza fronteiras de runs (caso menos comum, surge quando
  o documento foi editado manualmente em palavras-chave), o paragrafo eh
  reescrito com o texto novo no primeiro run com conteudo. Estilos parciais
  *dentro* do trecho substituido podem ser perdidos, mas a integridade do .docx
  fica preservada e o restante do paragrafo mantem sua formatacao.
"""

from __future__ import annotations

from io import BytesIO
from typing import Iterable

from docx import Document
from docx.text.paragraph import Paragraph


class SubstituicaoError(Exception):
    """Falha ao processar o documento (bytes invalidos, formato corrompido)."""


def extrair_texto(docx_bytes: bytes) -> str:
    """Extrai texto plano de um .docx, incluindo paragrafos do corpo e tabelas.

    O texto retornado eh enviado ao agente IA para identificar pares
    antigo<->novo de substituicao. Mantem ordem visual dos paragrafos para
    facilitar localizacao e validacao das ocorrencias.
    """
    if not docx_bytes:
        raise SubstituicaoError("Conteudo do documento vazio.")

    try:
        doc = Document(BytesIO(docx_bytes))
    except Exception as error:
        raise SubstituicaoError(
            f"Falha ao abrir o .docx: {type(error).__name__}: {error}"
        ) from error

    linhas: list[str] = []
    for paragraph in _iterar_paragrafos(doc):
        texto = paragraph.text
        if texto:
            linhas.append(texto)

    return "\n".join(linhas)


def aplicar_substituicoes(
    docx_bytes: bytes,
    pares: list[dict[str, str]],
) -> tuple[bytes, dict[str, int]]:
    """Aplica substituicoes em um .docx e retorna o novo binario + contagens.

    `pares`: lista de {"antigo": str, "novo": str}. Antigos vazios sao ignorados.
    Retorno: (bytes_editados, {antigo: ocorrencias_substituidas}).

    A contagem zero indica que o `antigo` nao foi encontrado — o caller decide
    se isso eh erro (ex.: validar contra ocorrencias_esperadas vindas do agente
    de IA). Esta funcao nao falha por ocorrencia ausente.
    """
    if not docx_bytes:
        raise SubstituicaoError("Conteudo do documento vazio.")

    try:
        doc = Document(BytesIO(docx_bytes))
    except Exception as error:
        raise SubstituicaoError(
            f"Falha ao abrir o .docx: {type(error).__name__}: {error}"
        ) from error

    pares_validos = [
        {"antigo": p["antigo"], "novo": p.get("novo", "")}
        for p in pares
        if p.get("antigo")
    ]
    ocorrencias: dict[str, int] = {p["antigo"]: 0 for p in pares_validos}

    for paragraph in _iterar_paragrafos(doc):
        for par in pares_validos:
            count = _substituir_em_paragrafo(paragraph, par["antigo"], par["novo"])
            ocorrencias[par["antigo"]] += count

    out = BytesIO()
    try:
        doc.save(out)
    except Exception as error:
        raise SubstituicaoError(
            f"Falha ao salvar o .docx editado: {type(error).__name__}: {error}"
        ) from error

    return out.getvalue(), ocorrencias


def _iterar_paragrafos(doc) -> Iterable[Paragraph]:
    """Itera todos os paragrafos do corpo + dentro de tabelas (recursivo)."""
    yield from doc.paragraphs
    for table in doc.tables:
        yield from _paragrafos_de_tabela(table)


def _paragrafos_de_tabela(table) -> Iterable[Paragraph]:
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from _paragrafos_de_tabela(nested)


def _substituir_em_paragrafo(paragraph: Paragraph, antigo: str, novo: str) -> int:
    """Substitui ocorrencias de `antigo` por `novo` em um paragrafo.

    Tenta primeiro substituir dentro de cada run individualmente (preserva estilo
    quando o trecho cabe num run so). Se ainda restar ocorrencia cruzando runs,
    aplica fallback de merge.
    """
    runs = paragraph.runs
    if not runs:
        return 0

    total = 0

    for run in runs:
        if antigo in run.text:
            n = run.text.count(antigo)
            run.text = run.text.replace(antigo, novo)
            total += n

    full_text = "".join(r.text for r in runs)
    if antigo not in full_text:
        return total

    cross = full_text.count(antigo)
    novo_text = full_text.replace(antigo, novo)
    primeiro_com_texto = next((r for r in runs if r.text), runs[0])
    primeiro_com_texto.text = novo_text
    for run in runs:
        if run is not primeiro_com_texto:
            run.text = ""
    return total + cross
